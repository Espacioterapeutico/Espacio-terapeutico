# -*- coding: utf-8 -*-
"""
Módulo de Autenticación, Seguridad y Administración General (routes_admin.py)
Encapsula Inicio de Sesión, Registro de Psicólogos, Gestión de Superadmin,
Restablecimiento de Contraseñas, Preguntas de Seguridad y Tickets de Soporte Técnico.
"""

import os
import json
import sqlite3
import datetime
from functools import wraps
from flask import Blueprint, request, jsonify, session, g, redirect
from werkzeug.security import generate_password_hash, check_password_hash

admin_bp = Blueprint('admin', __name__)

def get_db():
    """Obtiene la conexión a la base de datos desde el contexto global g de Flask."""
    db = getattr(g, '_database', None)
    if db is None:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(base_dir, 'clinica.db')
        db = g._database = sqlite3.connect(db_path, timeout=30.0)
        db.row_factory = sqlite3.Row
    return db

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': 'No autorizado. Por favor inicia sesión.'}), 401
        return f(*args, **kwargs)
    return decorated_function

def patient_login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'patient_id' not in session:
            return jsonify({'error': 'Sesión expirada o no iniciada.'}), 401
        return f(*args, **kwargs)
    return decorated_function

def check_is_superadmin():
    return session.get('role') in ['superadmin', 'admin']

# --- AUTENTICACIÓN Y SESIONES ---

@admin_bp.route('/api/register-admin', methods=['POST'])
def register_admin():
    data = request.json or {}
    username = data.get('username')
    password = data.get('password')
    
    if not username or not password:
        return jsonify({'error': 'Usuario y contraseña son requeridos.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    
    cursor.execute("SELECT id FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
    if cursor.fetchone():
        return jsonify({'error': 'El nombre de usuario ya está registrado.'}), 400
        
    password_hash = generate_password_hash(password)
    
    cursor.execute("SELECT COUNT(id) FROM usuarios")
    user_count = cursor.fetchone()[0] or 0
    user_role = 'superadmin' if user_count == 0 else 'psicologo'
    user_activo = 1 if user_role == 'superadmin' else 0
    
    try:
        cursor.execute("""
            INSERT INTO usuarios (username, password_hash, nombres, apellidos, role, activo)
            VALUES (?, ?, 'Administrador', 'General', ?, ?)
        """, (username, password_hash, user_role, user_activo))
        db.commit()
        return jsonify({'success': f'Usuario {user_role} creado con éxito.'})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al registrar administrador: {str(e)}'}), 500

@admin_bp.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.json or {}
        username = (data.get('username') or '').strip()
        password = (data.get('password') or '').strip()
        
        if not username or not password:
            return jsonify({'error': 'Usuario y contraseña son requeridos.'}), 400
            
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
        user = cursor.fetchone()
        
        if user and check_password_hash(user['password_hash'], password):
            u_dict = dict(user)
            
            # Verificar vencimiento de prueba gratis de 3 días para psicólogos no pagados
            if user['role'] == 'psicologo' and u_dict.get('suscripcion_paga', 0) != 1:
                expiry_str = u_dict.get('fecha_expiracion_prueba')
                if expiry_str:
                    try:
                        expiry_dt = datetime.datetime.strptime(expiry_str, "%Y-%m-%d %H:%M:%S")
                        if datetime.datetime.now() > expiry_dt:
                            cursor.execute("UPDATE usuarios SET activo = 0 WHERE id = ?", (user['id'],))
                            db.commit()
                            return jsonify({'error': 'Tu periodo de prueba gratis ha vencido. Contacta al administrador para activar tu suscripción.'}), 403
                    except Exception: pass
                        
            session.permanent = True
            session['user_id'] = user['id']
            session['username'] = user['username']
            session['role'] = user['role']
            session['activo'] = user['activo']
            u_dict = dict(user)
            return jsonify({
                'success': 'Inicio de sesión correcto.',
                'username': username,
                'nombres': u_dict.get('nombres') or '',
                'apellidos': u_dict.get('apellidos') or '',
                'role': user['role'],
                'activo': user['activo'],
                'aviso_pago': u_dict.get('aviso_pago', 0),
                'user_id': user['id'],
                'primer_inicio': u_dict.get('primer_inicio', 1) if u_dict.get('primer_inicio') is not None else 1,
                'suscripcion_paga': u_dict.get('suscripcion_paga', 0),
                'fecha_expiracion_prueba': u_dict.get('fecha_expiracion_prueba', ''),
                'bloqueos': {
                    'registro': u_dict.get('bloqueo_registro', 0),
                    'evoluciones': u_dict.get('bloqueo_evoluciones', 0),
                    'finanzas': u_dict.get('bloqueo_finanzas', 0),
                    'agenda': u_dict.get('bloqueo_agenda', 0),
                    'mensajes': u_dict.get('bloqueo_mensajes', 0),
                    'pizarra': 0 if user['role'] in ['psicologo', 'admin'] else (u_dict.get('bloqueo_pizarra') or 0),
                    'herramientas': 0 if user['role'] in ['psicologo', 'admin'] else (u_dict.get('bloqueo_herramientas') or 0),
                    'confirmaciones': 0 if user['role'] in ['psicologo', 'admin'] else (u_dict.get('bloqueo_confirmaciones') or 0),
                    'examen_mental': 0 if user['role'] in ['psicologo', 'admin'] else (u_dict.get('bloqueo_examen_mental') or 0),
                    'tests': 0 if user['role'] in ['psicologo', 'admin'] else (u_dict.get('bloqueo_tests') or 0)
                }
            })
        
        return jsonify({'error': 'Credenciales inválidas.'}), 401
    except Exception as e:
        return jsonify({'error': f'Error en el servidor al iniciar sesión: {str(e)}'}), 500

@admin_bp.route('/logout', methods=['GET', 'POST'])
@admin_bp.route('/api/logout', methods=['GET', 'POST'])
def logout():
    session.clear()
    session.modified = True
    if request.method == 'GET':
        return redirect('/login')
    return jsonify({'success': 'Sesión cerrada exitosamente.'})

@admin_bp.route('/api/check-username-role', methods=['GET'])
def check_username_role():
    username = request.args.get('username', '').strip()
    if not username:
        return jsonify({'error': 'Nombre de usuario requerido.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    
    cursor.execute("SELECT id FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
    if cursor.fetchone():
        return jsonify({'role': 'psicologo'})
        
    cursor.execute("SELECT id FROM pacientes WHERE LOWER(username) = ? OR cedula = ?", (username.lower(), username))
    if cursor.fetchone():
        return jsonify({'role': 'paciente'})
        
    return jsonify({'error': 'Usuario no encontrado.'}), 404

@admin_bp.route('/api/check-session', methods=['GET'])
def check_session():
    if 'user_id' in session:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SELECT * FROM usuarios WHERE id = ?", (session['user_id'],))
        row = cursor.fetchone()
        r_dict = dict(row) if row else {}

        return jsonify({
            'authenticated': True,
            'user_type': 'therapist',
            'user_id': session['user_id'],
            'username': session.get('username'),
            'nombres': r_dict.get('nombres') or '',
            'apellidos': r_dict.get('apellidos') or '',
            'role': r_dict.get('role', 'psicologo'),
            'activo': r_dict.get('activo', 1),
            'aviso_pago': r_dict.get('aviso_pago', 0),
            'primer_inicio': r_dict.get('primer_inicio', 1) if r_dict.get('primer_inicio') is not None else 1,
            'suscripcion_paga': r_dict.get('suscripcion_paga', 0),
            'fecha_expiracion_prueba': r_dict.get('fecha_expiracion_prueba', ''),
            'bloqueos': {
                'registro': r_dict.get('bloqueo_registro', 0),
                'evoluciones': r_dict.get('bloqueo_evoluciones', 0),
                'finanzas': r_dict.get('bloqueo_finanzas', 0),
                'agenda': r_dict.get('bloqueo_agenda', 0),
                'mensajes': r_dict.get('bloqueo_mensajes', 0),
                'pizarra': 0 if r_dict.get('role') in ['psicologo', 'admin'] else (r_dict.get('bloqueo_pizarra') or 0),
                'herramientas': 0 if r_dict.get('role') in ['psicologo', 'admin'] else (r_dict.get('bloqueo_herramientas') or 0),
                'confirmaciones': 0 if r_dict.get('role') in ['psicologo', 'admin'] else (r_dict.get('bloqueo_confirmaciones') or 0),
                'examen_mental': 0 if r_dict.get('role') in ['psicologo', 'admin'] else (r_dict.get('bloqueo_examen_mental') or 0),
                'tests': 0 if r_dict.get('role') in ['psicologo', 'admin'] else (r_dict.get('bloqueo_tests') or 0)
            }
        })
    elif 'patient_id' in session:
        return jsonify({
            'authenticated': True,
            'user_type': 'patient',
            'patient_id': session['patient_id'],
            'username': session.get('patient_username')
        })
    else:
        return jsonify({'authenticated': False}), 401

# --- RECUPERACIÓN Y CAMBIO DE CONTRASEÑAS ---

@admin_bp.route('/api/auth/get-security-questions', methods=['POST'])
def get_security_questions():
    data = request.json or {}
    username = (data.get('username') or '').strip()
    
    if not username:
        return jsonify({'error': 'Usuario es requerido.'}), 400
        
    db = get_db()
    cursor = db.cursor()

    cursor.execute("SELECT pregunta_seguridad_1, pregunta_seguridad_2 FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
    user_row = cursor.fetchone()

    if user_row and user_row['pregunta_seguridad_1'] and user_row['pregunta_seguridad_2']:
        return jsonify({
            'found': True,
            'pregunta_1': user_row['pregunta_seguridad_1'],
            'pregunta_2': user_row['pregunta_seguridad_2']
        })

    cursor.execute("SELECT pregunta_seguridad_1, pregunta_seguridad_2 FROM pacientes WHERE LOWER(username) = ? OR cedula = ?", (username.lower(), username))
    patient_row = cursor.fetchone()

    if patient_row and patient_row['pregunta_seguridad_1'] and patient_row['pregunta_seguridad_2']:
        return jsonify({
            'found': True,
            'pregunta_1': patient_row['pregunta_seguridad_1'],
            'pregunta_2': patient_row['pregunta_seguridad_2']
        })

    return jsonify({'error': 'El usuario no tiene configuradas preguntas de seguridad o no existe.'}), 404

@admin_bp.route('/api/admin/security-questions', methods=['GET', 'POST'])
@login_required
def admin_security_questions():
    db = get_db()
    cursor = db.cursor()
    user_id = session['user_id']

    if request.method == 'GET':
        cursor.execute("SELECT pregunta_seguridad_1, pregunta_seguridad_2 FROM usuarios WHERE id = ?", (user_id,))
        row = cursor.fetchone()
        return jsonify({
            'pregunta_seguridad_1': row['pregunta_seguridad_1'] if row else None,
            'pregunta_seguridad_2': row['pregunta_seguridad_2'] if row else None
        })

    data = request.json or {}
    p1 = (data.get('pregunta_seguridad_1') or '').strip()
    r1 = (data.get('respuesta_seguridad_1') or '').strip().lower()
    p2 = (data.get('pregunta_seguridad_2') or '').strip()
    r2 = (data.get('respuesta_seguridad_2') or '').strip().lower()

    if not p1 or not r1 or not p2 or not r2:
        return jsonify({'error': 'Las dos preguntas y respuestas son obligatorias.'}), 400

    r1_hash = generate_password_hash(r1)
    r2_hash = generate_password_hash(r2)

    cursor.execute("""
        UPDATE usuarios 
        SET pregunta_seguridad_1 = ?, respuesta_seguridad_1_hash = ?,
            pregunta_seguridad_2 = ?, respuesta_seguridad_2_hash = ?
        WHERE id = ?
    """, (p1, r1_hash, p2, r2_hash, user_id))
    db.commit()

    return jsonify({'success': 'Preguntas de seguridad de terapeuta guardadas con éxito.'})

@admin_bp.route('/api/user/change-password', methods=['POST'])
@login_required
def user_change_password():
    data = request.json or {}
    current_password = data.get('current_password')
    new_password = data.get('new_password')
    confirm_password = data.get('confirm_password')

    if not current_password or not new_password:
        return jsonify({'error': 'Ambas contraseñas son obligatorias.'}), 400

    if confirm_password and new_password != confirm_password:
        return jsonify({'error': 'La nueva contraseña y su confirmación no coinciden.'}), 400

    if len(new_password) < 6:
        return jsonify({'error': 'La nueva contraseña debe tener al menos 6 caracteres.'}), 400

    user_id = session['user_id']
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT password_hash FROM usuarios WHERE id = ?", (user_id,))
    user = cursor.fetchone()

    if not user or not check_password_hash(user['password_hash'], current_password):
        return jsonify({'error': 'La contraseña actual es incorrecta.'}), 401

    password_hash = generate_password_hash(new_password)
    cursor.execute("UPDATE usuarios SET password_hash = ? WHERE id = ?", (password_hash, user_id))
    db.commit()

    return jsonify({'success': 'Contraseña del terapeuta actualizada con éxito.'})

# --- SUPERADMIN & SOPORTE TÉCNICO ---

@admin_bp.route('/api/superadmin/therapists', methods=['GET'])
@login_required
def get_superadmin_therapists():
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
    db = get_db()
    cursor = db.cursor()
    cursor.execute("""
        SELECT u.*, 
               o.nombre as organizacion_nombre, 
               o.codigo_clinica, 
               o.modo_whatsapp, 
               o.admin_user_id as org_admin_user_id
        FROM usuarios u
        LEFT JOIN organizaciones o ON u.organizacion_id = o.id
        ORDER BY u.organizacion_id DESC, u.id DESC
    """)
    rows = cursor.fetchall()
    return jsonify([dict(r) for r in rows])

@admin_bp.route('/api/support/send', methods=['POST'])
@login_required
def send_support_ticket():
    data = request.json or {}
    asunto = data.get('asunto', '').strip()
    mensaje = data.get('mensaje', '').strip()

    if not asunto or not mensaje:
        return jsonify({'error': 'Asunto y mensaje son obligatorios.'}), 400

    db = get_db()
    cursor = db.cursor()
    user_id = session.get('user_id')
    now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    cursor.execute("""
        INSERT INTO soporte (user_id, asunto, mensaje, fecha, estado)
        VALUES (?, ?, ?, ?, 'Abierto')
    """, (user_id, asunto, mensaje, now_str))
    db.commit()

    return jsonify({'success': 'Tu mensaje de soporte ha sido enviado. El equipo técnico te contactará pronto.'})

def generate_default_slug_for_user(u):
    if not u:
        return ""
    if u['slug'] and u['slug'].strip():
        return u['slug'].strip()
    
    nom = (u['nombres'] or '').strip()
    ape = (u['apellidos'] or '').strip()
    uname = (u['username'] or '').strip()
    
    if nom or ape:
        combo = f"psic.{nom}{ape}"
    else:
        combo = f"psic.{uname}"
        
    import unicodedata, re
    normalized = unicodedata.normalize('NFD', combo)
    slug = re.sub(r'[\u0300-\u036f]', '', normalized).lower()
    slug = re.sub(r'[^a-z0-9\.]', '', slug)
    return slug or f"psic.{uname.lower()}"

DEFAULT_TERMS_TEXT = """Términos y Condiciones del Encuadre Terapéutico
Estimado/a consultante:

A continuación se presentan las condiciones operativas que rigen nuestro proceso terapéutico. Este marco tiene como objetivo proteger el tiempo de ambos, garantizar el compromiso mutuo y brindar una estructura clara a nuestras sesiones.

1. Duración de la Sesión
Cada sesión tiene una duración estimada de entre 45 minutos y 1 hora.

2. Gestión del Tiempo y Tardanzas
Retrasos por parte del consultante:
- Si llegas con retraso a la cita, el tiempo extra se otorgará únicamente si la agenda del terapeuta lo permite.
- Si el terapeuta tiene consultas posteriores, la sesión finalizará a la hora programada originalmente para no afectar el espacio de otros consultantes, aprovechando únicamente los minutos restantes.
- Tolerancia máxima: Pasados 15 minutos de retraso sin notificación, la consulta se considerará como consulta perdida (asistencia fallida) y deberá ser abonada en su totalidad (o descontada del paquete activo).

Retrasos por parte del terapeuta:
- En caso de que el terapeuta inicie la sesión con retraso, se garantizará el cumplimiento del tiempo total asignado (45 a 60 minutos), adaptando la agenda para no perjudicar al consultante.

3. Confirmación, Cancelación y Tiempo de Gracia
Confirmación de la cita:
- Toda sesión requiere confirmación previa. Si llega el día de la cita y esta no ha sido confirmada, el espacio no se reservará y la consulta se registrará automáticamente como cancelada con previo aviso.

Regla de cancelación y tiempo de gracia:
- Una vez confirmada la consulta, dispones de un tiempo de gracia de hasta 3 horas antes de la hora agendada para realizar cualquier cambio o cancelación sin costo alguno (cancelación con aviso).
- Pasado dicho límite (menos de 3 horas antes de la sesión), si cancelas o no te presentas, la sesión se computará como realizada y deberá ser abonada en su totalidad.

4. Paquetes de Sesiones
En caso de contar con un paquete de sesiones prepagado, cualquier inasistencia o cancelación fuera del tiempo de gracia permitido se descontará automáticamente del saldo de sesiones disponibles.

5. Cancelaciones por Parte del Terapeuta
Si el terapeuta debiera cancelar una sesión sin el debido aviso previo, asume el compromiso de reprogramar la consulta en la fecha disponible más próxima, habilitando de ser necesario fines de semana o días feriados para garantizar la atención oportuna.

Al agendar y confirmar tus sesiones a través de la plataforma, declaras haber leído y aceptado estos Términos y Condiciones para el desarrollo del proceso terapéutico."""

@admin_bp.route('/api/admin/profile-public', methods=['GET', 'POST'])
@login_required
def admin_profile_public():
    """Permite a cada psicólogo personalizar su perfil público (foto, biografía, modalidades, WhatsApp, redes, especialidades, ubicación)."""
    db = get_db()
    from app import ensure_usuarios_columns
    ensure_usuarios_columns(db)
    cursor = db.cursor()
    user_id = session['user_id']
    
    if request.method == 'GET':
        cursor.execute("""
            SELECT id, nombres, apellidos, username, slug, estudios, foto_titulo,
                   nomenclatura, descripcion_biografia, modalidades_json, whatsapp_publico, email_publico, redes_sociales_json,
                   especialidades, poblaciones_json, pais_ubicacion
            FROM usuarios WHERE id = ?
        """, (user_id,))
        u = cursor.fetchone()
        if not u:
            return jsonify({'error': 'Usuario no encontrado.'}), 404
            
        modalidades_raw = u['modalidades_json']
        modalidades_data = {}
        modalidades_list = []
        if modalidades_raw:
            try:
                parsed = json.loads(modalidades_raw)
                if isinstance(parsed, dict):
                    modalidades_data = parsed
                    if parsed.get('online'): modalidades_list.append("Online")
                    if parsed.get('presencial'): modalidades_list.append("Presencial")
                    if parsed.get('domicilio'): modalidades_list.append("Domicilio")
                elif isinstance(parsed, list):
                    modalidades_list = parsed
                    modalidades_data = {
                        'online': 'Online' in parsed,
                        'online_titulo': 'Consulta Online',
                        'online_detalle': '',
                        'presencial': 'Presencial' in parsed,
                        'presencial_titulo': 'Consulta Presencial',
                        'presencial_detalle': '',
                        'domicilio': 'Domicilio' in parsed,
                        'domicilio_titulo': 'Atención a Domicilio',
                        'domicilio_detalle': ''
                    }
            except Exception:
                modalidades_list = ["Online", "Presencial"]
                modalidades_data = {'online': True, 'online_titulo': 'Consulta Online', 'presencial': True, 'presencial_titulo': 'Consulta Presencial'}
        else:
            modalidades_list = ["Online", "Presencial"]
            modalidades_data = {'online': True, 'online_titulo': 'Consulta Online', 'presencial': True, 'presencial_titulo': 'Consulta Presencial'}

        redes = json.loads(u['redes_sociales_json']) if u['redes_sociales_json'] else {}
        poblaciones = json.loads(u['poblaciones_json']) if u['poblaciones_json'] else ["Adultos", "Adolescentes"]
        
        slug = u['slug'] or generate_default_slug_for_user(u)
        clean_slug = slug.replace('psic.', '') if slug.startswith('psic.') else slug

        return jsonify({
            'slug': slug,
            'clean_slug': clean_slug,
            'full_profile_url': f"/psic.{clean_slug}",
            'registration_url': f"/registro/psic.{clean_slug}",
            'fast_booking_url': f"/agendar/psic.{clean_slug}",
            'nomenclatura': u['nomenclatura'] or u['estudios'] or 'Psicólogo Clínico',
            'descripcion_biografia': u['descripcion_biografia'] or '',
            'especialidades': u['especialidades'] or '',
            'pais_ubicacion': u['pais_ubicacion'] or '',
            'pais': u['pais_ubicacion'] or '',
            'poblaciones': poblaciones,
            'modalidades': modalidades_list,
            'modalidades_data': modalidades_data,
            'whatsapp_publico': u['whatsapp_publico'] or '',
            'email_publico': u['email_publico'] or '',
            'redes_sociales': redes,
            'foto': u['foto_titulo'] or '/static/logo.png'
        })
    else:
        data = request.json or {}
        nomenclatura = data.get('nomenclatura', '').strip()
        descripcion = data.get('descripcion_biografia', '').strip()
        especialidades = data.get('especialidades', '').strip()
        pais_ubicacion = data.get('pais_ubicacion') or data.get('pais', '')
        poblaciones = data.get('poblaciones', ["Adultos", "Adolescentes"])
        
        mods_data = data.get('modalidades_data')
        if not mods_data:
            mods_data = data.get('modalidades', ["Online", "Presencial"])

        whatsapp = data.get('whatsapp_publico', '').strip()
        email = data.get('email_publico', '').strip()
        redes = data.get('redes_sociales', {})
        foto = data.get('foto', '')
        
        try:
            cursor.execute("""
                UPDATE usuarios SET 
                    nomenclatura = ?,
                    descripcion_biografia = ?,
                    especialidades = ?,
                    pais_ubicacion = ?,
                    poblaciones_json = ?,
                    modalidades_json = ?,
                    whatsapp_publico = ?,
                    email_publico = ?,
                    redes_sociales_json = ?,
                    foto_titulo = CASE WHEN ? != '' THEN ? ELSE foto_titulo END
                WHERE id = ?
            """, (nomenclatura, descripcion, especialidades, pais_ubicacion, json.dumps(poblaciones), json.dumps(mods_data), whatsapp, email, json.dumps(redes), foto, foto, user_id))
            db.commit()
            return jsonify({'success': 'Perfil público actualizado con éxito.'})
        except Exception as e:
            return jsonify({'error': f'Error al actualizar perfil: {str(e)}'}), 500

@admin_bp.route('/api/admin/payment-methods', methods=['GET', 'POST'])
@login_required
def admin_payment_methods():
    """Obtiene o actualiza las instrucciones de pago del psicólogo."""
    db = get_db()
    cursor = db.cursor()
    user_id = session.get('user_id')
    
    if request.method == 'GET':
        cursor.execute("SELECT metodos_pago FROM usuarios WHERE id = ?", (user_id,))
        row = cursor.fetchone()
        metodos = row['metodos_pago'] if row and row['metodos_pago'] else ""
        return jsonify({'metodos_pago': metodos})
    else:
        data = request.json or {}
        metodos = data.get('metodos_pago', '')
        cursor.execute("UPDATE usuarios SET metodos_pago = ? WHERE id = ?", (metodos, user_id))
        db.commit()
        return jsonify({'success': 'Métodos de pago actualizados con éxito.'})

@admin_bp.route('/api/admin/terms', methods=['GET', 'POST'])
@login_required
def admin_terms():
    """Obtiene o actualiza los Términos y Condiciones del psicólogo."""
    db = get_db()
    cursor = db.cursor()
    user_id = session.get('user_id')
    
    if request.method == 'GET':
        cursor.execute("SELECT terminos_condiciones FROM usuarios WHERE id = ?", (user_id,))
        row = cursor.fetchone()
        terms = (row['terminos_condiciones'] if row and row['terminos_condiciones'] else "").strip()
        if not terms:
            terms = DEFAULT_TERMS_TEXT
        return jsonify({'terms': terms})
    else:
        data = request.json or {}
        terms = data.get('terms', '').strip()
        cursor.execute("UPDATE usuarios SET terminos_condiciones = ? WHERE id = ?", (terms, user_id))
        db.commit()
        return jsonify({'success': 'Términos y Condiciones actualizados con éxito.'})

# --- GOOGLE CALENDAR ENDPOINTS ---
CLIENT_SECRETS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "credentials.json")
SCOPES = ['https://www.googleapis.com/auth/calendar']

try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import Flow
    from googleapiclient.discovery import build
    from google.auth.transport.requests import Request
    GOOGLE_CALENDAR_AVAILABLE = True
except ImportError:
    GOOGLE_CALENDAR_AVAILABLE = False

def get_calendar_service(user_id=None):
    if not GOOGLE_CALENDAR_AVAILABLE:
        return None
    db = get_db()
    cursor = db.cursor()
    
    if not user_id:
        try:
            user_id = session.get('user_id')
        except RuntimeError:
            user_id = None
            
    if not user_id:
        return None
            
    token_key = f'google_token_{user_id}'
    cursor.execute("SELECT valor FROM configuracion WHERE clave = ?", (token_key,))
    row = cursor.fetchone()
    
    if not row:
        return None
        
    try:
        creds_data = json.loads(row['valor'])
        creds = Credentials.from_authorized_user_info(creds_data, SCOPES)
        
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
            cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES (?, ?)", 
                           (token_key, creds.to_json()))
            db.commit()
            
        return build('calendar', 'v3', credentials=creds)
    except Exception as e:
        print("Error al inicializar servicio de Google Calendar:", e)
        return None

@admin_bp.route('/api/google/status', methods=['GET'])
@login_required
def google_status():
    try:
        has_credentials_json = os.path.exists(CLIENT_SECRETS_FILE)
        service = get_calendar_service()
        return jsonify({
            'configured': service is not None,
            'has_credentials_json': has_credentials_json
        })
    except Exception as e:
        return jsonify({
            'configured': False,
            'has_credentials_json': os.path.exists(CLIENT_SECRETS_FILE),
            'error': str(e)
        }), 200

@admin_bp.route('/api/google/upload-credentials', methods=['POST'])
@login_required
def upload_google_credentials():
    if 'file' not in request.files:
        return jsonify({'error': 'No se proporcionó ningún archivo.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío.'}), 400
    if not file.filename.endswith('.json'):
        return jsonify({'error': 'El archivo debe ser en formato JSON.'}), 400
    try:
        content = file.read().decode('utf-8')
        config_data = json.loads(content)
        if 'web' not in config_data and 'installed' not in config_data:
            return jsonify({'error': 'El archivo no es un JSON de credenciales de Google válido.'}), 400
        
        with open(CLIENT_SECRETS_FILE, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, indent=4)
            
        return jsonify({'success': 'Credenciales subidas e instaladas con éxito.'})
    except Exception as e:
        return jsonify({'error': f'Error al procesar el archivo: {str(e)}'}), 500

@admin_bp.route('/api/superadmin/stats', methods=['GET'])
@login_required
def superadmin_get_stats():
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    
    try:
        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        cursor.execute("SELECT COUNT(id) FROM pacientes")
        total_pacientes = cursor.fetchone()[0] or 0
        
        cursor.execute("SELECT COUNT(id) FROM usuarios WHERE role = 'psicologo'")
        total_psicologos = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(id) FROM usuarios 
            WHERE role = 'psicologo' 
              AND COALESCE(activo, 1) = 1 
              AND (COALESCE(suscripcion_paga, 0) = 1 OR (fecha_expiracion_prueba IS NOT NULL AND fecha_expiracion_prueba > ?))
        """, (now_str,))
        psicologos_activos = cursor.fetchone()[0] or 0
        
        cursor.execute("""
            SELECT COUNT(id) FROM usuarios 
            WHERE role = 'psicologo' 
              AND (
                COALESCE(activo, 1) = 0 
                OR (COALESCE(suscripcion_paga, 0) = 0 AND (fecha_expiracion_prueba IS NULL OR fecha_expiracion_prueba <= ?))
              )
        """, (now_str,))
        psicologos_deudores = cursor.fetchone()[0] or 0
        
        return jsonify({
            'total_pacientes': total_pacientes,
            'total_psicologos': total_psicologos,
            'psicologos_activos': psicologos_activos,
            'psicologos_deudores': psicologos_deudores
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@admin_bp.route('/api/superadmin/therapists/<int:user_id>/update-profile', methods=['POST'])
@login_required
def superadmin_update_therapist_profile(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    try:
        data = request.json or {}
        db = get_db()
        from app import ensure_usuarios_columns
        ensure_usuarios_columns(db)
        cursor = db.cursor()
        
        cursor.execute("SELECT id FROM usuarios WHERE id = ?", (user_id,))
        if not cursor.fetchone():
            return jsonify({'error': 'Psicólogo no encontrado.'}), 404
            
        nombres = (data.get('nombres') or '').strip()
        apellidos = (data.get('apellidos') or '').strip()
        username = (data.get('username') or '').strip()
        email = (data.get('email') or '').strip()
        email_publico = (data.get('email_publico') or '').strip()
        cedula = (data.get('cedula') or '').strip()
        whatsapp = (data.get('whatsapp_publico') or '').strip()
        bio = (data.get('descripcion_biografia') or '').strip()

        if username:
            cursor.execute("SELECT id FROM usuarios WHERE LOWER(username) = LOWER(?) AND id != ?", (username, user_id))
            if cursor.fetchone():
                return jsonify({'error': f"El nombre de usuario '@{username}' ya está registrado en la plataforma."}), 400

        cursor.execute("PRAGMA table_info(usuarios)")
        existing_cols = {row[1] for row in cursor.fetchall()}

        fields_to_update = {}
        if 'nombres' in existing_cols: fields_to_update['nombres'] = nombres
        if 'apellidos' in existing_cols: fields_to_update['apellidos'] = apellidos
        if 'username' in existing_cols and username: fields_to_update['username'] = username
        if 'email' in existing_cols: fields_to_update['email'] = email
        if 'email_publico' in existing_cols: fields_to_update['email_publico'] = email_publico
        if 'cedula' in existing_cols: fields_to_update['cedula'] = cedula
        if 'whatsapp_publico' in existing_cols: fields_to_update['whatsapp_publico'] = whatsapp
        if 'telefono' in existing_cols and 'whatsapp_publico' not in existing_cols: fields_to_update['telefono'] = whatsapp
        if 'descripcion_biografia' in existing_cols: fields_to_update['descripcion_biografia'] = bio
        if 'slug' in existing_cols and username: fields_to_update['slug'] = f"psic.{username.lower()}"

        if fields_to_update:
            set_clause = ", ".join([f"{col} = ?" for col in fields_to_update.keys()])
            params = list(fields_to_update.values()) + [user_id]
            cursor.execute(f"UPDATE usuarios SET {set_clause} WHERE id = ?", params)
            db.commit()

        return jsonify({'success': 'Ficha del psicólogo actualizada con éxito.'})
    except Exception as ex:
        print("[SUPERADMIN] Error actualizando ficha de psicólogo:", ex)
        return jsonify({'error': f"Error en base de datos: {str(ex)}"}), 500

@admin_bp.route('/api/superadmin/therapists/<int:user_id>/save-settings', methods=['POST'])
@login_required
def superadmin_save_therapist_settings(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    data = request.json or {}
    db = get_db()
    from app import ensure_usuarios_columns
    ensure_usuarios_columns(db)
    cursor = db.cursor()
    
    mostrar_en_directorio = 1 if data.get('mostrar_en_directorio') else 0
    aviso_pago = 1 if data.get('aviso_pago') else 0
    bloqueo_registro = 1 if data.get('bloqueo_registro') else 0
    bloqueo_evoluciones = 1 if data.get('bloqueo_evoluciones') else 0
    bloqueo_finanzas = 1 if data.get('bloqueo_finanzas') else 0
    bloqueo_agenda = 1 if data.get('bloqueo_agenda') else 0
    bloqueo_mensajes = 1 if data.get('bloqueo_mensajes') else 0
    bloqueo_pizarra = 1 if data.get('bloqueo_pizarra') else 0
    bloqueo_herramientas = 1 if data.get('bloqueo_herramientas') else 0
    bloqueo_confirmaciones = 1 if data.get('bloqueo_confirmaciones') else 0
    
    cursor.execute("""
        UPDATE usuarios 
        SET mostrar_en_directorio = ?, aviso_pago = ?,
            bloqueo_registro = ?, bloqueo_evoluciones = ?, bloqueo_finanzas = ?,
            bloqueo_agenda = ?, bloqueo_mensajes = ?, bloqueo_pizarra = ?, bloqueo_herramientas = ?, bloqueo_confirmaciones = ?
        WHERE id = ?
    """, (mostrar_en_directorio, aviso_pago, bloqueo_registro, bloqueo_evoluciones, bloqueo_finanzas,
          bloqueo_agenda, bloqueo_mensajes, bloqueo_pizarra, bloqueo_herramientas, bloqueo_confirmaciones, user_id))
    db.commit()
    return jsonify({'success': '¡Cambios guardados con éxito en la base de datos!'})

@admin_bp.route('/api/superadmin/therapists/<int:user_id>/toggle-feature', methods=['POST'])
@login_required
def superadmin_toggle_feature(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    data = request.json or {}
    feature = data.get('feature')
    status = data.get('status')
    
    if feature not in ['registro', 'evoluciones', 'finanzas', 'agenda', 'mensajes', 'pizarra', 'herramientas']:
        return jsonify({'error': 'Función no válida.'}), 400
        
    if status not in [0, 1]:
        return jsonify({'error': 'Estado de bloqueo no válido.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    column = f"bloqueo_{feature}"
    cursor.execute(f"UPDATE usuarios SET {column} = ? WHERE id = ? AND role = 'psicologo'", (status, user_id))
    db.commit()
    
    return jsonify({'success': f'Función {feature} actualizada con éxito.', 'feature': feature, 'status': status})


# --- RUTAS MIGRADAS AUTOMÁTICAMENTE DE AUDITORÍA ---

@admin_bp.route('/api/admin-exists', methods=['GET'])
def admin_exists():
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id FROM usuarios LIMIT 1")
    user = cursor.fetchone()
    return jsonify({'exists': user is not None})



@admin_bp.route('/api/register/check-cedula', methods=['GET'])
def check_register_cedula():
    cedula = request.args.get('cedula', '').strip()
    if not cedula:
        return jsonify({'error': 'Cédula es requerida.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id, username, nombres, apellidos, cedula, pregunta_seguridad_1, respuesta_seguridad_1_hash FROM pacientes")
    rows = cursor.fetchall()
    
    cleaned_input = ''.join(c for c in cedula if c.isdigit())
    row = None
    for r in rows:
        db_cedula = r['cedula'] or ''
        if db_cedula.strip() == cedula:
            row = r
            break
        if cleaned_input and ''.join(c for c in db_cedula if c.isdigit()) == cleaned_input:
            row = r
            break
            
    if row:
        if row['pregunta_seguridad_1'] and row['respuesta_seguridad_1_hash']:
            return jsonify({'status': 'registered'})
        else:
            return jsonify({
                'status': 'pre_registered', 
                'nombres': row['nombres'], 
                'apellidos': row['apellidos']
            })
            
    return jsonify({'status': 'new_patient'})



@admin_bp.route('/api/register', methods=['POST'])
def register():
    data = request.json
    tipo_usuario = data.get('tipo_usuario') # 'psicologo' o 'paciente'
    username = data.get('username')
    password = data.get('password')
    
    if not username or not password or not tipo_usuario:
        return jsonify({'error': 'Usuario, contraseña y tipo de usuario son requeridos.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    
    # Verificar si el usuario ya existe en alguna de las dos tablas
    cursor.execute("SELECT id FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
    if cursor.fetchone():
        return jsonify({'error': 'El nombre de usuario ya está registrado.'}), 400
        
    cursor.execute("SELECT id FROM pacientes WHERE LOWER(username) = ?", (username.lower(),))
    if cursor.fetchone():
        return jsonify({'error': 'El nombre de usuario ya está registrado.'}), 400

    password_hash = generate_password_hash(password)
    
    try:
        if tipo_usuario == 'psicologo':
            nombres = data.get('nombres')
            apellidos = data.get('apellidos')
            estudios = data.get('estudios')
            federacion = data.get('federacion')
            foto_titulo = data.get('foto_titulo', '')
            foto_documento = data.get('foto_documento', '')
            
            import datetime
            now_dt = datetime.datetime.now()
            expiry_dt = now_dt + datetime.timedelta(days=30)
            now_str = now_dt.strftime("%Y-%m-%d %H:%M:%S")
            expiry_str = expiry_dt.strftime("%Y-%m-%d %H:%M:%S")

            import unicodedata
            clean_name = f"psic.{(nombres or '').strip()}{(apellidos or '').strip()}".lower().replace(" ", "")
            if len(clean_name) <= 5:
                clean_name = f"psic.{username.lower()}"
            clean_slug = re.sub(r'[^a-z0-9\.]', '', unicodedata.normalize('NFD', clean_name))

            default_visual_cfg = json.dumps({
                "duracion": 60,
                "costo_online": 30.0,
                "costo_presencial": 35.0,
                "moneda": "USD",
                "alerta_confirmacion": 2,
                "perfiles": [
                    {
                        "nombre": "Horario Estándar",
                        "activo": True,
                        "dias": [
                            {"dia": 1, "nombre": "Lunes", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}, {"inicio": "14:00", "fin": "18:00"}]},
                            {"dia": 2, "nombre": "Martes", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}, {"inicio": "14:00", "fin": "18:00"}]},
                            {"dia": 3, "nombre": "Miércoles", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}, {"inicio": "14:00", "fin": "18:00"}]},
                            {"dia": 4, "nombre": "Jueves", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}, {"inicio": "14:00", "fin": "18:00"}]},
                            {"dia": 5, "nombre": "Viernes", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}, {"inicio": "14:00", "fin": "18:00"}]}
                        ]
                    }
                ]
            })

            default_pm_str = "Pago Móvil / Transferencia Bancaria\nZelle / PayPal disponible"

            p1 = data.get('pregunta_seguridad_1')
            r1 = (data.get('respuesta_seguridad_1') or '').strip().lower()
            p2 = data.get('pregunta_seguridad_2')
            r2 = (data.get('respuesta_seguridad_2') or '').strip().lower()
            r1_hash = generate_password_hash(r1) if r1 else None
            r2_hash = generate_password_hash(r2) if r2 else None

            cursor.execute("""
                INSERT INTO usuarios (username, password_hash, nombres, apellidos, estudios, federacion, foto_titulo, foto_documento, role, activo, fecha_registro, fecha_expiracion_prueba, suscripcion_paga, slug, configuracion_horarios_visual, metodos_pago, primer_inicio, pregunta_seguridad_1, respuesta_seguridad_1_hash, pregunta_seguridad_2, respuesta_seguridad_2_hash)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'psicologo', 1, ?, ?, 0, ?, ?, ?, 1, ?, ?, ?, ?)
            """, (username, password_hash, nombres, apellidos, estudios, federacion, foto_titulo, foto_documento, now_str, expiry_str, clean_slug, default_visual_cfg, default_pm_str, p1, r1_hash, p2, r2_hash))
            db.commit()
            
            # Enviar correo de bienvenida con credenciales y preguntas de seguridad
            email_target = username if '@' in username else data.get('email')
            if email_target:
                full_name_psic = f"Psic. {nombres} {apellidos}".strip()
                send_welcome_credentials_email('psicologo', email_target, full_name_psic, username, password, p1, r1, p2, r2)

            return jsonify({'success': 'Cuenta de psicólogo creada con éxito. Tienes 1 mes (30 días) de prueba gratuita.'})
            
        elif tipo_usuario == 'paciente':
            nombres = data.get('nombres')
            apellidos = data.get('apellidos')
            cedula = data.get('cedula')
            telefono = data.get('telefono')
            email = data.get('email')
            pronombre = data.get('pronombre')
            genero = data.get('genero')
            edad = data.get('edad')
            lugar_nacimiento = data.get('lugar_nacimiento')
            fecha_nacimiento = data.get('fecha_nacimiento')
            residencia_actual = data.get('residencia_actual')
            pais = data.get('pais')
            ciudad = data.get('ciudad')
            con_quien_reside = data.get('con_quien_reside')
            nivel_academico = data.get('nivel_academico')
            ocupacion = data.get('ocupacion')
            estado_civil = data.get('estado_civil')
            contacto_emergencia_nombre = data.get('contacto_emergencia_nombre')
            contacto_emergencia_parentesco = data.get('contacto_emergencia_parentesco')
            motivo_consulta = data.get('motivo_consulta')
            expectativas = data.get('expectativas')
            farmacologia = data.get('farmacologia')
            pregunta_1 = data.get('pregunta_seguridad_1')
            resp_1 = data.get('respuesta_seguridad_1')
            pregunta_2 = data.get('pregunta_seguridad_2')
            resp_2 = data.get('respuesta_seguridad_2')
            psicologo_id = data.get('psicologo_id')
            
            # Verificar si la cédula ya existe (comparación flexible)
            cursor.execute("SELECT id, username, cedula, nombres, apellidos, psicologo_id, pregunta_seguridad_1, respuesta_seguridad_1_hash FROM pacientes")
            all_patients = cursor.fetchall()
            cleaned_input = ''.join(c for c in cedula if c.isdigit()) if cedula else ''
            existing_patient = None
            for p in all_patients:
                db_cedula = p['cedula'] or ''
                if cedula and db_cedula.strip() == cedula:
                    existing_patient = p
                    break
                if cleaned_input and ''.join(c for c in db_cedula if c.isdigit()) == cleaned_input:
                    existing_patient = p
                    break
            
            resp_1_hash = generate_password_hash(resp_1) if resp_1 else None
            resp_2_hash = generate_password_hash(resp_2) if resp_2 else None
            
            if existing_patient:
                if existing_patient['pregunta_seguridad_1'] and existing_patient['respuesta_seguridad_1_hash']:
                    return jsonify({'error': 'La cédula ya está registrada con una cuenta activa.'}), 400
                
                # Paciente pre-registrado: actualizar credenciales de acceso y campos mínimos
                cursor.execute("""
                    UPDATE pacientes
                    SET username = ?, password_hash = ?,
                        pregunta_seguridad_1 = ?, respuesta_seguridad_1_hash = ?,
                        pregunta_seguridad_2 = ?, respuesta_seguridad_2_hash = ?,
                        telefono = COALESCE(?, telefono),
                        email = COALESCE(?, email)
                    WHERE id = ?
                """, (
                    username, password_hash, 
                    pregunta_1, resp_1_hash, 
                    pregunta_2, resp_2_hash,
                    telefono, email,
                    existing_patient['id']
                ))
                patient_id = existing_patient['id']
                target_psic = psicologo_id or existing_patient['psicologo_id'] or 1
                if psicologo_id and not existing_patient['psicologo_id']:
                    cursor.execute("UPDATE pacientes SET psicologo_id = ? WHERE id = ?", (psicologo_id, patient_id))
                ex_nom = existing_patient['nombres'] if existing_patient['nombres'] else ''
                ex_ape = existing_patient['apellidos'] if existing_patient['apellidos'] else ''
                pat_name = f"{data.get('nombres') or ex_nom} {data.get('apellidos') or ex_ape}".strip() or username
                notif_msg = f"El consultante {pat_name} ha completado la creación de su cuenta de acceso."
            else:
                target_psic = psicologo_id or 1
                cursor.execute("""
                    INSERT INTO pacientes (
                        nombres, apellidos, cedula, telefono, email, pronombre, genero, edad,
                        lugar_nacimiento, fecha_nacimiento, residencia_actual, pais, ciudad, con_quien_reside,
                        nivel_academico, ocupacion, estado_civil, contacto_emergencia_nombre,
                        contacto_emergencia_parentesco, motivo_consulta, expectativas, farmacologia,
                        username, password_hash, pregunta_seguridad_1, respuesta_seguridad_1_hash,
                        pregunta_seguridad_2, respuesta_seguridad_2_hash, psicologo_id
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    nombres, apellidos, cedula, telefono, email, pronombre, genero, edad,
                    lugar_nacimiento, fecha_nacimiento, residencia_actual, pais, ciudad, con_quien_reside,
                    nivel_academico, ocupacion, estado_civil, contacto_emergencia_nombre,
                    contacto_emergencia_parentesco, motivo_consulta, expectativas, farmacologia,
                    username, password_hash, pregunta_1, resp_1_hash, pregunta_2, resp_2_hash, target_psic
                ))
                patient_id = cursor.lastrowid
                pat_name = f"{nombres} {apellidos}".strip() or username
                notif_msg = f"El consultante {pat_name} se ha registrado en la plataforma."

            # Generar notificación interna y push al psicólogo asignado
            from datetime import datetime
            now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute("""
                INSERT INTO notificaciones (user_id, tipo, titulo, mensaje, fecha, leida, link)
                VALUES (?, 'nuevo_paciente', '👤 Nuevo Registro de Consultante', ?, ?, 0, '/#pacientes')
            """, (target_psic, notif_msg, now_str))
            
            send_fcm_notification(user_id=target_psic, title="👤 Nuevo Registro de Consultante", body=notif_msg, url="/#pacientes")
            send_webpush_notification(user_id=target_psic, title="👤 Nuevo Registro de Consultante", body=notif_msg, url="/#pacientes")

            db.commit()
            
            # Enviar correo de bienvenida con credenciales y preguntas de seguridad al consultante
            if email:
                full_name_pac = pat_name
                send_welcome_credentials_email(
                    'paciente', email, full_name_pac, username, password, 
                    pregunta_1, resp_1, pregunta_2, resp_2, 
                    login_url="https://www.espacioterapeutico.net/portal"
                )
            
            # Sincronización en segundo plano con Firebase
            import threading
            threading.Thread(target=sync_patient_to_firebase, args=(patient_id,)).start()
            
            return jsonify({'success': 'Cuenta de consultante creada con éxito.'})
            
        else:
            return jsonify({'error': 'Tipo de usuario no válido.'}), 400
            
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al registrar: {str(e)}'}), 500



@admin_bp.route('/api/superadmin/seed-demo-user', methods=['GET', 'POST'])
def superadmin_seed_demo_user():
    try:
        db = get_db()
        ensure_demo_user(db)
        return jsonify({'success': True, 'message': 'Usuario demo psicologa.valeria y paciente camila.perez procesados con éxito.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@admin_bp.route('/api/superadmin/create-psychologist', methods=['POST'])
@login_required
def superadmin_create_psychologist():
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    data = request.json
    nombres = data.get('nombres')
    apellidos = data.get('apellidos')
    username = data.get('username')
    password = data.get('password')
    estudios = data.get('estudios')
    cedula = data.get('cedula', '')
    federacion = data.get('federacion')
    foto_titulo = data.get('foto_titulo', '')
    foto_documento = data.get('foto_documento', '')
    
    if not username or not password or not nombres or not apellidos:
        return jsonify({'error': 'Todos los campos requeridos deben ser completados.'}), 400
        
    db = get_db()
    ensure_usuarios_columns(db)
    cursor = db.cursor()
    
    cursor.execute("SELECT id FROM usuarios WHERE LOWER(username) = ?", (username.lower(),))
    if cursor.fetchone():
        return jsonify({'error': 'El nombre de usuario ya existe.'}), 400
        
    password_hash = generate_password_hash(password)
    
    try:
        import datetime
        now_dt = datetime.datetime.now()
        expiry_dt = now_dt + datetime.timedelta(days=30)
        now_str = now_dt.strftime("%Y-%m-%d %H:%M:%S")
        expiry_str = expiry_dt.strftime("%Y-%m-%d %H:%M:%S")

        clean_name = f"psic.{(nombres or '').strip()}{(apellidos or '').strip()}".lower().replace(" ", "")
        if not clean_name or clean_name == "psic.":
            clean_slug = f"psic.{username.lower()}"
        else:
            import unicodedata, re
            clean_slug = re.sub(r'[^a-z0-9\.]', '', unicodedata.normalize('NFD', clean_name))

        cursor.execute("""
            INSERT INTO usuarios (username, password_hash, nombres, apellidos, cedula, estudios, federacion, foto_titulo, foto_documento, role, activo, fecha_registro, fecha_expiracion_prueba, suscripcion_paga, slug)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'psicologo', 1, ?, ?, 0, ?)
        """, (username, password_hash, nombres, apellidos, cedula, estudios, federacion, foto_titulo, foto_documento, now_str, expiry_str, clean_slug))
        db.commit()
        return jsonify({'success': 'Psicólogo registrado con éxito (Modo Prueba 1 Mes / 30 Días activo).'})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al registrar psicólogo: {str(e)}'}), 500



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/toggle-active', methods=['POST'])
@login_required
def superadmin_toggle_active(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT activo FROM usuarios WHERE id = ? AND role = 'psicologo'", (user_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    new_status = 0 if row['activo'] == 1 else 1
    if new_status == 1:
        import datetime
        now_dt = datetime.datetime.now()
        expiry_dt = now_dt + datetime.timedelta(days=30)
        now_str = now_dt.strftime("%Y-%m-%d %H:%M:%S")
        expiry_str = expiry_dt.strftime("%Y-%m-%d %H:%M:%S")
        cursor.execute("UPDATE usuarios SET activo = 1, fecha_registro = COALESCE(fecha_registro, ?), fecha_expiracion_prueba = ? WHERE id = ?", (now_str, expiry_str, user_id))
    else:
        cursor.execute("UPDATE usuarios SET activo = 0 WHERE id = ?", (user_id,))
    db.commit()
    return jsonify({'success': 'Estado de suscripción actualizado.', 'activo': new_status})



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/toggle-subscription', methods=['POST'])
@login_required
def superadmin_toggle_subscription(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT suscripcion_paga FROM usuarios WHERE id = ? AND role = 'psicologo'", (user_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    new_status = 0 if row['suscripcion_paga'] == 1 else 1
    cursor.execute("UPDATE usuarios SET suscripcion_paga = ? WHERE id = ?", (new_status, user_id))
    db.commit()
    return jsonify({'success': 'Estado de suscripción paga actualizado.', 'suscripcion_paga': new_status})



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/set-expiration', methods=['POST'])
@login_required
def superadmin_set_therapist_expiration(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    data = request.json or {}
    fecha_exp = data.get('fecha_expiracion') or data.get('fecha_expiracion_prueba')
    suscripcion_paga = 1 if data.get('suscripcion_paga') else 0
    
    if suscripcion_paga == 0:
        # Si se desactiva, forzar fecha de expiración a ayer a las 23:59:59
        yesterday_dt = datetime.datetime.now() - datetime.timedelta(days=1)
        fecha_exp = yesterday_dt.strftime('%Y-%m-%d 23:59:59')
    elif not fecha_exp:
        return jsonify({'error': 'Fecha de expiración/renovación es requerida.'}), 400
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id FROM usuarios WHERE id = ?", (user_id,))
    if not cursor.fetchone():
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    if len(str(fecha_exp)) == 10:
        fecha_exp = f"{fecha_exp} 23:59:59"
        
    cursor.execute("""
        UPDATE usuarios 
        SET fecha_expiracion_prueba = ?, suscripcion_paga = ? 
        WHERE id = ?
    """, (fecha_exp, suscripcion_paga, user_id))

    # Si el usuario pertenece a una clínica y es Director Admin o se solicitó aplicar a la clínica completa:
    cursor.execute("SELECT organizacion_id, tipo_clinica FROM usuarios WHERE id = ?", (user_id,))
    u_info = cursor.fetchone()
    if u_info and u_info['organizacion_id'] and (u_info['tipo_clinica'] == 1 or data.get('apply_to_clinic')):
        cursor.execute("""
            UPDATE usuarios 
            SET fecha_expiracion_prueba = ?, suscripcion_paga = ? 
            WHERE organizacion_id = ?
        """, (fecha_exp, suscripcion_paga, u_info['organizacion_id']))

    db.commit()

    email_status_msg = ""
    # Enviar correo de notificación de activación/renovación solo si la suscripción se marca como activa
    if suscripcion_paga == 1:
        try:
            cursor.execute("SELECT nombres, apellidos, username, email, email_publico FROM usuarios WHERE id = ?", (user_id,))
            usr = cursor.fetchone()
            if usr:
                u_dict = dict(usr)
                target_email = (u_dict.get('email') or u_dict.get('email_publico') or '').strip()
                if target_email:
                    full_name = f"{u_dict.get('nombres') or ''} {u_dict.get('apellidos') or ''}".strip() or u_dict.get('username')
                    try:
                        exp_dt = datetime.datetime.strptime(str(fecha_exp)[:10], '%Y-%m-%d')
                        exp_formatted = exp_dt.strftime('%d/%m/%Y')
                    except Exception:
                        exp_formatted = str(fecha_exp)[:10]
                    send_subscription_renewed_email(target_email, full_name, exp_formatted)
                    email_status_msg = f"📧 Correo de renovación enviado a {target_email}"
                else:
                    email_status_msg = "⚠️ El psicólogo no tiene correo electrónico guardado. Usa '📋 Ficha' para agregarlo."
        except Exception as ex_mail:
            print("[SMTP] Error enviando correo de renovación:", ex_mail)
            email_status_msg = f"⚠️ Ocurrió una duda al enviar correo: {ex_mail}"

    res_text = '⛔ Suscripción desactivada inmediatamente.' if suscripcion_paga == 0 else 'Fecha de expiración / renovación actualizada con éxito.'
    return jsonify({'success': res_text, 'email_status': email_status_msg})



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/toggle-aviso-pago', methods=['POST'])
@login_required
def superadmin_toggle_aviso_pago(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT aviso_pago FROM usuarios WHERE id = ? AND role = 'psicologo'", (user_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    new_status = 0 if row['aviso_pago'] == 1 else 1
    cursor.execute("UPDATE usuarios SET aviso_pago = ? WHERE id = ?", (new_status, user_id))
    db.commit()
    return jsonify({'success': True, 'aviso_pago': new_status})



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/toggle-directorio', methods=['POST'])
@login_required
def superadmin_toggle_directorio(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT COALESCE(mostrar_en_directorio, 1) as mostrar_en_directorio FROM usuarios WHERE id = ?", (user_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    new_val = 0 if row['mostrar_en_directorio'] == 1 else 1
    cursor.execute("UPDATE usuarios SET mostrar_en_directorio = ? WHERE id = ?", (new_val, user_id))
    db.commit()
    return jsonify({'success': 'Visibilidad en directorio actualizada.', 'mostrar_en_directorio': new_val})



@admin_bp.route('/api/superadmin/therapists/<int:user_id>/update-documents', methods=['POST'])
@login_required
def superadmin_update_therapist_documents(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    data = request.json or {}
    foto_titulo = data.get('foto_titulo')
    foto_documento = data.get('foto_documento')
    
    db = get_db()
    cursor = db.cursor()
    if foto_titulo is not None:
        cursor.execute("UPDATE usuarios SET foto_titulo = ? WHERE id = ?", (foto_titulo, user_id))
    if foto_documento is not None:
        cursor.execute("UPDATE usuarios SET foto_documento = ? WHERE id = ?", (foto_documento, user_id))
    db.commit()
    return jsonify({'success': 'Documento actualizado con éxito.'})



@admin_bp.route('/api/superadmin/support', methods=['GET'])
@login_required
def superadmin_get_support_tickets():
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id, usuario_id, rol_remitente, nombre_remitente, email_remitente, mensaje, fecha, leido FROM soporte ORDER BY id DESC")
    rows = cursor.fetchall()
    return jsonify([dict(r) for r in rows])



@admin_bp.route('/api/superadmin/support/<int:ticket_id>/mark-read', methods=['POST'])
@login_required
def superadmin_mark_ticket_read(ticket_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("UPDATE soporte SET leido = 1 WHERE id = ?", (ticket_id,))
    db.commit()
    return jsonify({'success': 'Ticket marcado como leído.'})



@admin_bp.route('/api/superadmin/support/<int:ticket_id>', methods=['DELETE'])
@login_required
def superadmin_delete_ticket(ticket_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("DELETE FROM soporte WHERE id = ?", (ticket_id,))
    db.commit()
    return jsonify({'success': 'Ticket eliminado.'})



@admin_bp.route('/api/sync/auto-backup', methods=['POST', 'GET'])
def auto_backup():
    path = create_automatic_backup()
    return jsonify({'success': True, 'backup': path})



@admin_bp.route('/api/cron/hourly-tool-reminders', methods=['GET', 'POST'])
def cron_hourly_tool_reminders():
    try:
        count = send_hourly_patient_tool_reminders()
        return jsonify({'success': True, 'reminders_sent': count})
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@admin_bp.route('/api/auth/reset-password', methods=['POST'])
def auth_reset_password():
    data = request.json or {}
    username = (data.get('username') or '').strip()
    respuesta_1 = (data.get('respuesta_1') or '').strip().lower()
    respuesta_2 = (data.get('respuesta_2') or '').strip().lower()
    new_password = (data.get('new_password') or '').strip()
    
    if not username or not respuesta_1 or not respuesta_2 or not new_password:
        return jsonify({'error': 'Todos los campos son obligatorios.'}), 400

    if len(new_password) < 6:
        return jsonify({'error': 'La nueva contraseña debe tener al menos 6 caracteres.'}), 400

    db = get_db()
    cursor = db.cursor()

    # 1. Intentar en usuarios (psicólogos/admin)
    cursor.execute("""
        SELECT id, respuesta_seguridad_1_hash, respuesta_seguridad_2_hash 
        FROM usuarios 
        WHERE LOWER(username) = ?
    """, (username.lower(),))
    user_row = cursor.fetchone()

    if user_row and user_row['respuesta_seguridad_1_hash'] and user_row['respuesta_seguridad_2_hash']:
        match_1 = check_password_hash(user_row['respuesta_seguridad_1_hash'], respuesta_1)
        match_2 = check_password_hash(user_row['respuesta_seguridad_2_hash'], respuesta_2)
        if match_1 and match_2:
            new_hash = generate_password_hash(new_password)
            cursor.execute("UPDATE usuarios SET password_hash = ? WHERE id = ?", (new_hash, user_row['id']))
            db.commit()
            return jsonify({'success': 'Contraseña restablecida con éxito. Ya puedes iniciar sesión.'})
        else:
            return jsonify({'error': 'Respuestas a preguntas de seguridad incorrectas.'}), 401

    # 2. Intentar en pacientes
    cursor.execute("""
        SELECT id, respuesta_seguridad_1_hash, respuesta_seguridad_2_hash 
        FROM pacientes 
        WHERE LOWER(username) = ? OR cedula = ?
    """, (username.lower(), username))
    patient_row = cursor.fetchone()

    if patient_row and patient_row['respuesta_seguridad_1_hash'] and patient_row['respuesta_seguridad_2_hash']:
        match_1 = check_password_hash(patient_row['respuesta_seguridad_1_hash'], respuesta_1)
        match_2 = check_password_hash(patient_row['respuesta_seguridad_2_hash'], respuesta_2)
        if match_1 and match_2:
            new_hash = generate_password_hash(new_password)
            cursor.execute("UPDATE pacientes SET password_hash = ? WHERE id = ?", (new_hash, patient_row['id']))
            db.commit()
            try:
                import threading
                threading.Thread(target=sync_patient_to_firebase, args=(patient_row['id'],)).start()
            except Exception:
                pass
            return jsonify({'success': 'Contraseña restablecida con éxito. Ya puedes iniciar sesión.'})
        else:
            return jsonify({'error': 'Respuestas a preguntas de seguridad incorrectas.'}), 401

    return jsonify({'error': 'El usuario no existe o no tiene preguntas de seguridad configuradas.'}), 404



@admin_bp.route('/api/admin/reset-test-data', methods=['POST'])
@login_required
def reset_test_data():
    db = get_db()
    cursor = db.cursor()
    try:
        tables_to_clear = [
            'agenda_finanzas',
            'sesiones',
            'pizarra_terapeutica',
            'notificaciones',
            'soporte',
            'pagos_notificados'
        ]
        for table in tables_to_clear:
            cursor.execute(f"DELETE FROM {table}")
            cursor.execute(f"DELETE FROM sqlite_sequence WHERE name='{table}'")
            
        cursor.execute("UPDATE pacientes SET psicologo_id = 1 WHERE cedula IN ('26540973', '84586641')")
        db.commit()
        
        # Sincronizar Firebase si corresponde
        try:
            sync_patient_to_firebase(3) # Leo
            sync_patient_to_firebase(7) # Eulogio
        except:
            pass
            
        return jsonify({'success': 'Datos de consultas, evoluciones y pagos restablecidos a cero con éxito.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@admin_bp.route('/api/admin/smtp-settings', methods=['GET', 'POST'])
@login_required
def admin_smtp_settings():
    db = get_db()
    cursor = db.cursor()
    
    keys = ['smtp_host', 'smtp_port', 'smtp_user', 'smtp_password', 'smtp_from_email', 'auto_welcome_email_active']
    
    if request.method == 'GET':
        settings = {}
        for key in keys:
            cursor.execute("SELECT valor FROM configuracion WHERE clave = ?", (key,))
            row = cursor.fetchone()
            settings[key] = row['valor'] if row else ""
        if not settings.get('smtp_host'):
            settings['smtp_host'] = 'smtp.gmail.com'
        if not settings.get('smtp_port'):
            settings['smtp_port'] = '587'
        if not settings.get('smtp_user'):
            settings['smtp_user'] = 'espacioterapeuticoapp@gmail.com'
        if not settings.get('smtp_password'):
            settings['smtp_password'] = 'kinygwxtkovrtsjp'
        if not settings.get('smtp_from_email'):
            settings['smtp_from_email'] = 'Espacio Terapéutico <espacioterapeuticoapp@gmail.com>'
        if not settings.get('auto_welcome_email_active'):
            settings['auto_welcome_email_active'] = '1'
        return jsonify(settings)
        
    data = request.json or {}
    try:
        for key in keys:
            if key in data:
                cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES (?, ?)", (key, str(data[key]).strip()))
        db.commit()
        return jsonify({'success': 'Configuración de servidor SMTP actualizada con éxito.'})
    except Exception as e:
        return jsonify({'error': f'Error al guardar configuración SMTP: {str(e)}'}), 500



@admin_bp.route('/api/admin/smtp-test', methods=['POST'])
@login_required
def admin_smtp_test():
    data = request.json or {}
    test_email = data.get('test_email', '').strip()
    if not test_email or '@' not in test_email:
        return jsonify({'error': 'Proporciona una dirección de correo válida para la prueba.'}), 400

    db = get_db()
    cursor = db.cursor()
    cursor.execute("""
        SELECT clave, valor FROM configuracion 
        WHERE clave IN ('smtp_host', 'smtp_port', 'smtp_user', 'smtp_password', 'smtp_from_email')
    """)
    cfg = {r['clave']: r['valor'] for r in cursor.fetchall()}
    
    smtp_host = data.get('smtp_host') or cfg.get('smtp_host', '').strip()
    smtp_port_raw = str(data.get('smtp_port') or cfg.get('smtp_port', '587')).strip()
    smtp_port = int(smtp_port_raw) if smtp_port_raw.isdigit() else 587
    smtp_user = data.get('smtp_user') or cfg.get('smtp_user', '').strip()
    smtp_pass = data.get('smtp_password') or cfg.get('smtp_password', '').strip()
    smtp_from = data.get('smtp_from_email') or cfg.get('smtp_from_email', '').strip() or f"Espacio Terapéutico <{smtp_user}>"

    if not smtp_host or not smtp_user or not smtp_pass:
        return jsonify({'error': 'Incompleto. Servidor Host, Usuario y Contraseña SMTP son requeridos.'}), 400

    try:
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText

        msg = MIMEMultipart('alternative')
        msg['Subject'] = "🧪 Prueba de Conexión SMTP - Espacio Terapéutico"
        msg['From'] = smtp_from
        msg['To'] = test_email

        html_body = f"""
        <div style="font-family: Arial, sans-serif; padding: 20px; background: #f4f6f9;">
            <div style="max-width: 500px; margin: 0 auto; background: #fff; padding: 24px; border-radius: 12px; border: 1px solid #10b981;">
                <h2 style="color: #10b981; margin-top: 0;">✅ Servidor de Correo Operativo</h2>
                <p>¡Hola! Este es un correo de prueba enviado exitosamente desde <strong>Espacio Terapéutico</strong>.</p>
                <p><strong>Configuración Verificada:</strong></p>
                <ul>
                    <li>Servidor: <code>{smtp_host}:{smtp_port}</code></li>
                    <li>Usuario Remitente: <code>{smtp_user}</code></li>
                </ul>
                <p style="font-size: 12px; color: #64748b;">Tu plataforma ya está lista para enviar credenciales y notificaciones a psicólogos y consultantes.</p>
            </div>
        </div>
        """
        msg.attach(MIMEText(html_body, 'html', 'utf-8'))

        if smtp_port == 465:
            server = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=12)
        else:
            server = smtplib.SMTP(smtp_host, smtp_port, timeout=12)
            server.starttls()

        server.login(smtp_user, smtp_pass)
        server.sendmail(smtp_user, [test_email], msg.as_string())
        server.quit()

        return jsonify({'success': f'¡Correo de prueba enviado con éxito a {test_email}!'})
    except Exception as e:
        return jsonify({'error': f'Falló la conexión SMTP: {str(e)}'}), 500



@admin_bp.route('/api/admin/profile-slug', methods=['GET', 'POST'])
@login_required
def admin_profile_slug():
    db = get_db()
    cursor = db.cursor()
    user_id = session['user_id']
    if request.method == 'GET':
        cursor.execute("SELECT id, nombres, apellidos, username, slug FROM usuarios WHERE id = ?", (user_id,))
        u = cursor.fetchone()
        if not u:
            return jsonify({'error': 'Usuario no encontrado.'}), 404
        
        current_slug = generate_default_slug_for_user(u)
        if not u['slug']:
            try:
                cursor.execute("UPDATE usuarios SET slug = ? WHERE id = ?", (current_slug, user_id))
                db.commit()
            except Exception:
                pass

        return jsonify({
            'id': u['id'],
            'username': u['username'],
            'slug': current_slug,
            'fast_booking_url': f"/agendar/{current_slug}",
            'registration_url': f"/registro/{current_slug}"
        })
    else:
        data = request.json or {}
        new_slug = str(data.get('slug', '')).strip().lower().replace(" ", "-")
        new_slug = re.sub(r'[^a-z0-9\.\-_]', '', new_slug)
        if not new_slug:
            return jsonify({'error': 'El identificador personalizado (slug) no puede estar vacío.'}), 400
        
        cursor.execute("SELECT id FROM usuarios WHERE (LOWER(slug) = ? OR LOWER(username) = ?) AND id != ?", (new_slug, new_slug, user_id))
        if cursor.fetchone():
            return jsonify({'error': 'El enlace personalizado ya está en uso por otro profesional.'}), 400
            
        cursor.execute("UPDATE usuarios SET slug = ? WHERE id = ?", (new_slug, user_id))
        db.commit()
        return jsonify({
            'success': 'Enlace personalizado actualizado con éxito.',
            'slug': new_slug,
            'fast_booking_url': f"/agendar/{new_slug}",
            'registration_url': f"/registro/{new_slug}"
        })



def ensure_usuarios_columns(db):
    try:
        cursor = db.cursor()
        cursor.execute("PRAGMA table_info(usuarios)")
        columns = [column[1] for column in cursor.fetchall()]
        needed_columns = {
            'especialidades': 'TEXT',
            'pais_ubicacion': 'TEXT',
            'poblaciones_json': 'TEXT',
            'nomenclatura': 'TEXT',
            'descripcion_biografia': 'TEXT',
            'modalidades_json': 'TEXT',
            'redes_sociales_json': 'TEXT',
            'terminos_condiciones': 'TEXT',
            'whatsapp_publico': 'TEXT',
            'email_publico': 'TEXT'
        }
        for col_name, col_type in needed_columns.items():
            if col_name not in columns:
                cursor.execute(f"ALTER TABLE usuarios ADD COLUMN {col_name} {col_type}")
        db.commit()
    except Exception as e:
        print("Aviso ensure_usuarios_columns:", e)

@admin_bp.route('/api/public/landing-content', methods=['GET'])
def get_public_landing_content():
    """Obtiene el contenido editable de la portada principal y el directorio público de psicólogos."""
    db = get_db()
    ensure_usuarios_columns(db)
    cursor = db.cursor()
    
    # 1. Textos institucionales editables por el Superadmin
    cursor.execute("SELECT clave, valor FROM configuracion WHERE clave LIKE 'landing_%'")
    cfg_rows = cursor.fetchall()
    content = {r['clave']: r['valor'] for r in cfg_rows}
    
    # 2. Directorio de psicólogos activos (excluye cuenta de sistema 'admin')
    cursor.execute("""
        SELECT id, nombres, apellidos, username, slug, estudios, foto_titulo, foto_documento,
               nomenclatura, descripcion_biografia, modalidades_json, whatsapp_publico, email_publico, redes_sociales_json
        FROM usuarios 
        WHERE (COALESCE(activo, 1) = 1) 
          AND (COALESCE(mostrar_en_directorio, 1) = 1) 
          AND (COALESCE(suscripcion_paga, 0) = 1 OR role = 'superadmin')
          AND (role IS NULL OR role = '' OR role = 'psicologo' OR role = 'admin' OR role = 'superadmin')
          AND LOWER(username) NOT IN ('admin', 'superadmin')
        ORDER BY id ASC
    """)
    therapists_rows = cursor.fetchall()
    therapists = []
    for t in therapists_rows:
        slug = generate_default_slug_for_user(t)
        modalidades = json.loads(t['modalidades_json']) if t['modalidades_json'] else ["Online", "Presencial"]
        redes = json.loads(t['redes_sociales_json']) if t['redes_sociales_json'] else {}
        therapists.append({
            'id': t['id'],
            'nombres': t['nombres'] or 'Psicólogo',
            'apellidos': t['apellidos'] or '',
            'nombre_completo': f"Psic. {t['nombres'] or ''} {t['apellidos'] or ''}".strip(),
            'slug': slug,
            'nomenclatura': t['nomenclatura'] or t['estudios'] or 'Psicólogo Clínico',
            'descripcion': t['descripcion_biografia'] or '',
            'foto': t['foto_titulo'] or '/static/logo.png',
            'modalidades': modalidades,
            'whatsapp': t['whatsapp_publico'] or '',
            'email': t['email_publico'] or '',
            'redes': redes,
            'url_perfil': f"/psic.{slug.replace('psic.', '') if slug.startswith('psic.') else slug}",
            'url_agendar': f"/agendar/{slug}",
            'url_registro': f"/registro/{slug}"
        })
        
    resp = jsonify({
        'content': content,
        'therapists': therapists
    })
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return resp



@admin_bp.route('/api/public/therapist/<path:slug>', methods=['GET'])
def get_public_therapist_profile(slug):
    db = get_db()
    ensure_usuarios_columns(db)
    cursor = db.cursor()
    from app import get_psychologist_by_id_or_slug
    
    clean_slug = str(slug).strip()
    for prefix in ['psicologo/', 'psic/', 'psic.', 'agendar/', 'registro/']:
        if clean_slug.lower().startswith(prefix):
            clean_slug = clean_slug[len(prefix):]
            
    psych = get_psychologist_by_id_or_slug(cursor, clean_slug)
    if not psych:
        # Re-try with original raw slug just in case
        psych = get_psychologist_by_id_or_slug(cursor, slug)
        
    if not psych:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404

    psych = dict(psych) if not isinstance(psych, dict) else psych

    modalidades_raw = psych.get('modalidades_json')
    modalidades_data = {}
    modalidades_list = []
    
    if modalidades_raw:
        try:
            parsed = json.loads(modalidades_raw)
            if isinstance(parsed, dict):
                modalidades_data = parsed
                if parsed.get('online'): modalidades_list.append("Online")
                if parsed.get('presencial'): modalidades_list.append("Presencial")
                if parsed.get('domicilio'): modalidades_list.append("Domicilio")
            elif isinstance(parsed, list):
                modalidades_list = parsed
                for m in parsed:
                    modalidades_data[str(m).lower()] = True
        except Exception:
            modalidades_list = ["Online", "Presencial"]
            modalidades_data = {'online': True, 'presencial': True}
    else:
        modalidades_list = ["Online", "Presencial"]
        modalidades_data = {'online': True, 'presencial': True}

    try:
        redes_raw = psych.get('redes_sociales_json')
        redes = json.loads(redes_raw) if redes_raw else {}
    except Exception:
        redes = {}

    try:
        poblaciones_raw = psych.get('poblaciones_json')
        poblaciones = json.loads(poblaciones_raw) if poblaciones_raw else []
    except Exception:
        poblaciones = []
        
    clean_slug = psych.get('slug') or generate_default_slug_for_user(psych)
    foto_url = psych.get('foto_perfil') or psych.get('foto_titulo') or '/static/logo.png'
    
    resp = jsonify({
        'id': psych.get('id'),
        'nombres': psych.get('nombres') or '',
        'apellidos': psych.get('apellidos') or '',
        'nombre_completo': f"Psic. {psych.get('nombres') or ''} {psych.get('apellidos') or ''}".strip(),
        'slug': clean_slug,
        'nomenclatura': psych.get('nomenclatura') or psych.get('estudios') or 'Psicólogo Clínico',
        'descripcion_biografia': psych.get('descripcion_biografia') or '',
        'especialidades': psych.get('especialidades') or '',
        'pais': psych.get('pais_ubicacion') or '',
        'poblaciones': poblaciones,
        'foto': foto_url,
        'modalidades': modalidades_list,
        'modalidades_data': modalidades_data,
        'whatsapp_publico': psych.get('whatsapp_publico') or '',
        'email_publico': psych.get('email_publico') or '',
        'redes_sociales': redes,
        'url_agendar': f"/agendar/{clean_slug}",
        'url_registro': f"/registro/{clean_slug}"
    })
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return resp



@admin_bp.route('/api/admin/landing-content', methods=['POST'])
@login_required
def update_admin_landing_content():
    """Permite al Superadmin actualizar los textos dinámicos de la portada web."""
    if not check_is_superadmin():
        return jsonify({'error': 'No tienes permisos para modificar la portada web.'}), 403
        
    data = request.json or {}
    db = get_db()
    cursor = db.cursor()
    
    allowed_keys = [
        'landing_hero_title', 'landing_hero_subtitle',
        'landing_quienes_somos', 'landing_mision', 'landing_vision',
        'landing_footer_text'
    ]
    
    try:
        for k in allowed_keys:
            if k in data:
                cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES (?, ?)", (k, str(data[k])))
        db.commit()
        return jsonify({'success': 'Contenidos de la portada web actualizados exitosamente.'})
    except Exception as e:
        return jsonify({'error': f'Error al actualizar contenidos: {str(e)}'}), 500



@admin_bp.route('/api/admin/rates', methods=['POST'])
@login_required
def admin_rates():
    data = request.json
    db = get_db()
    cursor = db.cursor()
    import json
    
    # Obtener configuración actual
    cursor.execute("SELECT configuracion_horarios_visual FROM usuarios WHERE id = ?", (session.get('user_id'),))
    u_row = cursor.fetchone()
    config = {}
    if u_row and u_row[0]:
        try:
            config = json.loads(u_row[0])
        except:
            pass
            
    config['tarifas'] = data.get('tarifas', {})
    config['paquetes'] = data.get('paquetes', {})
    
    try:
        cursor.execute("UPDATE usuarios SET configuracion_horarios_visual = ? WHERE id = ?", (json.dumps(config), session.get('user_id')))
        db.commit()
        return jsonify({'success': 'Tarifas y honorarios actualizados con éxito.'})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al actualizar tarifas: {str(e)}'}), 500

# ==========================================
# GESTIÓN DE PACIENTES
# ==========================================



@admin_bp.route('/api/admin/patients/<int:patient_id>/rates', methods=['PUT'])
@login_required
def update_patient_rates_quick(patient_id):
    try:
        data = request.json or {}
        db = get_db()
        cursor = db.cursor()
        psicologo_id = session.get('user_id')
        
        # Verificar que el paciente pertenece a este psicólogo
        cursor.execute("SELECT id FROM pacientes WHERE id = ? AND psicologo_id = ?", (patient_id, psicologo_id))
        if not cursor.fetchone():
            return jsonify({'error': 'Paciente no encontrado o sin acceso.'}), 404
        
        costo_personalizado = data.get('costo_personalizado')
        moneda_personalizada = data.get('moneda_personalizada')
        costo_paquete_personalizado = data.get('costo_paquete_personalizado')
        sesiones_paquete_personalizado = data.get('sesiones_paquete_personalizado')
        
        # Convertir vacíos a None
        if costo_personalizado == '' or costo_personalizado is None:
            costo_personalizado = None
        if costo_paquete_personalizado == '' or costo_paquete_personalizado is None:
            costo_paquete_personalizado = None
        if sesiones_paquete_personalizado == '' or sesiones_paquete_personalizado is None:
            sesiones_paquete_personalizado = None
        
        cursor.execute("""
            UPDATE pacientes SET
                costo_personalizado = ?,
                moneda_personalizada = ?,
                costo_paquete_personalizado = ?,
                sesiones_paquete_personalizado = ?
            WHERE id = ?
        """, (costo_personalizado, moneda_personalizada, costo_paquete_personalizado, sesiones_paquete_personalizado, patient_id))
        db.commit()
        
        import threading
        threading.Thread(target=sync_patient_to_firebase, args=(patient_id,)).start()
        
        return jsonify({'success': 'Honorarios actualizados con éxito.'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500



@admin_bp.route('/api/finance/transactions', methods=['POST'])
@login_required
def add_transaction():
    data = request.json
    db = get_db()
    cursor = db.cursor()
    
    paciente_id = data.get('paciente_id')
    fecha = data.get('fecha')
    hora = data.get('hora')
    tipo_consulta = data.get('tipo_consulta') # 'Presencial', 'Online'
    monto = data.get('monto', 0.0)
    moneda = data.get('moneda') # 'USD', 'EUR', 'BSD'
    estado_pago = data.get('estado_pago') # 'Paga', 'Pendiente', 'Prepagada'
    control_uso = data.get('control_uso', 'Consumida') # 'Consumida', 'No consumida'
    fecha_liquidacion = data.get('fecha_liquidacion')
    
    cantidad_sesiones = int(data.get('cantidad_sesiones', 1) or 1)
    referencia = data.get('referencia')
    metodo_pago = data.get('metodo_pago')
    fecha_pago = data.get('fecha_pago')
    
    if (estado_pago == 'Prepagada' or 'paquete' in (tipo_consulta or '').lower()) and cantidad_sesiones <= 1:
        cursor.execute("SELECT costo_paquete_personalizado, sesiones_paquete_personalizado FROM pacientes WHERE id = ?", (paciente_id,))
        pac = cursor.fetchone()
        if pac and pac['sesiones_paquete_personalizado']:
            cantidad_sesiones = int(pac['sesiones_paquete_personalizado'])
    
    if not paciente_id or not fecha or not tipo_consulta or not moneda or not estado_pago:
        return jsonify({'error': 'Faltan campos requeridos para la transacción.'}), 400
        
    try:
        cursor.execute("""
            INSERT INTO agenda_finanzas (
                paciente_id, fecha, hora, tipo_consulta, monto, moneda, 
                estado_pago, control_uso, fecha_liquidacion, cantidad_sesiones,
                referencia, metodo_pago, fecha_pago
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            paciente_id, fecha, hora, tipo_consulta, monto, moneda,
            estado_pago, control_uso, fecha_liquidacion, cantidad_sesiones,
            referencia, metodo_pago, fecha_pago
        ))
        db.commit()
        auto_settle_patient_debts(db, paciente_id)
        
        # Sincronización en segundo plano con Firebase
        import threading
        threading.Thread(target=sync_patient_to_firebase, args=(paciente_id,)).start()
        
        return jsonify({'success': 'Transacción agregada con éxito.'})
    except Exception as e:
        return jsonify({'error': f'Error al agregar transacción: {str(e)}'}), 500



@admin_bp.route('/api/finance/export-csv', methods=['GET'])
@login_required
def export_finance_csv():
    try:
        import io
        import csv
        from flask import Response
        
        month = request.args.get('month')
        year = request.args.get('year')
        
        if not month or not year:
            now = get_now_vet()
            month = f"{now.month:02d}"
            year = str(now.year)
        else:
            month = f"{int(month):02d}"
            year = str(year)
            
        date_prefix = f"{year}-{month}%"
        
        db = get_db()
        cursor = db.cursor()
        
        cursor.execute("""
            SELECT af.id, af.fecha, af.hora, p.nombres, p.apellidos, p.cedula,
                   af.tipo_consulta, af.monto, af.moneda, af.estado_pago,
                   af.control_uso, af.metodo_pago, af.referencia, af.fecha_liquidacion
            FROM agenda_finanzas af
            LEFT JOIN pacientes p ON af.paciente_id = p.id
            WHERE af.fecha LIKE ? OR af.fecha_liquidacion LIKE ?
            ORDER BY af.fecha DESC, af.hora DESC
        """, (date_prefix, date_prefix))
        
        rows = cursor.fetchall()
        
        output = io.StringIO()
        writer = csv.writer(output, delimiter=';')
        
        writer.writerow(['ID', 'Fecha Cita', 'Hora Cita', 'Consultante', 'Cedula', 'Modalidad', 'Monto', 'Moneda', 'Estado de Pago', 'Control Uso', 'Metodo Pago', 'Referencia', 'Fecha Liquidacion'])
        
        for r in rows:
            nombre_paciente = f"{r['nombres']} {r['apellidos']}" if r['nombres'] else "Consultante Desconocido"
            writer.writerow([
                r['id'],
                r['fecha'] or '',
                r['hora'] or '',
                nombre_paciente,
                r['cedula'] or '',
                r['tipo_consulta'] or '',
                f"{float(r['monto'] or 0):.2f}",
                r['moneda'] or 'USD',
                r['estado_pago'] or '',
                r['control_uso'] or '',
                r['metodo_pago'] or '',
                r['referencia'] or '',
                r['fecha_liquidacion'] or ''
            ])
            
        csv_data = output.getvalue()
        output.close()
        
        filename = f"Reporte_Financiero_{year}_{month}.csv"
        return Response(
            csv_data,
            mimetype="text/csv",
            headers={"Content-disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        return jsonify({'error': f'Error al exportar CSV: {str(e)}'}), 500



@admin_bp.route('/api/finance/transactions/<int:trans_id>', methods=['GET'])
@login_required
def get_transaction(trans_id):
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT * FROM agenda_finanzas WHERE id = ?", (trans_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Transacción no encontrada.'}), 404
    return jsonify(dict(row))



@admin_bp.route('/api/finance/transactions/<int:trans_id>', methods=['PUT'])
@login_required
def update_transaction(trans_id):
    data = request.json
    db = get_db()
    cursor = db.cursor()
    try:
        cursor.execute("SELECT * FROM agenda_finanzas WHERE id = ?", (trans_id,))
        row = cursor.fetchone()
        if not row:
            return jsonify({'error': 'Transacción no encontrada.'}), 404
            
        original_estado_pago = row['estado_pago']
        estado_pago = data.get('estado_pago') if 'estado_pago' in data else row['estado_pago']
        if original_estado_pago == 'Cancelada sin aviso' and estado_pago == 'Paga':
            estado_pago = 'Cancelada sin aviso - Paga'
        control_uso = data.get('control_uso') if 'control_uso' in data else row['control_uso']
        fecha_liquidacion = data.get('fecha_liquidacion') if 'fecha_liquidacion' in data else row['fecha_liquidacion']
        monto = data.get('monto') if 'monto' in data else row['monto']
        moneda = data.get('moneda') if 'moneda' in data else row['moneda']
        cantidad_sesiones = data.get('cantidad_sesiones') if 'cantidad_sesiones' in data else row['cantidad_sesiones']
        referencia = data.get('referencia') if 'referencia' in data else row['referencia']
        metodo_pago = data.get('metodo_pago') if 'metodo_pago' in data else row['metodo_pago']
        fecha_pago = data.get('fecha_pago') if 'fecha_pago' in data else row['fecha_pago']
        fecha = data.get('fecha') if 'fecha' in data else row['fecha']
        hora = data.get('hora') if 'hora' in data else row['hora']
        tipo_consulta = data.get('tipo_consulta') if 'tipo_consulta' in data else row['tipo_consulta']
        
        if estado_pago == 'ConsumirPrepago':
            cursor.execute("""
                SELECT id, cantidad_sesiones, control_uso 
                FROM agenda_finanzas 
                WHERE paciente_id = ? AND estado_pago = 'Prepagada' AND control_uso = 'No consumida'
                ORDER BY fecha ASC, id ASC LIMIT 1
            """, (row['paciente_id'],))
            pkg = cursor.fetchone()
            if not pkg:
                return jsonify({'error': 'El consultante no tiene sesiones prepagadas disponibles.'}), 400
                
            pkg_id = pkg['id']
            pkg_cant = pkg['cantidad_sesiones']
            if pkg_cant > 1:
                cursor.execute("UPDATE agenda_finanzas SET cantidad_sesiones = ? WHERE id = ?", (pkg_cant - 1, pkg_id))
            else:
                cursor.execute("UPDATE agenda_finanzas SET control_uso = 'Consumida' WHERE id = ?", (pkg_id,))
                
            estado_pago = 'Paga'
            monto = 0.0
            cantidad_sesiones = 1
            control_uso = 'Consumida'
        
        # Sincronizar actualización con Google Calendar si está enlazado
        google_event_id = row['google_event_id']
        if google_event_id:
            service = get_calendar_service()
            if service:
                start_datetime = f"{fecha}T{hora}:00"
                end_hour = str(int(hora.split(':')[0]) + 1).zfill(2)
                end_datetime = f"{fecha}T{end_hour}:{hora.split(':')[1]}:00"
                try:
                    g_event = service.events().get(calendarId='primary', eventId=google_event_id).execute()
                    g_event['start'] = {'dateTime': start_datetime, 'timeZone': 'America/Caracas'}
                    g_event['end'] = {'dateTime': end_datetime, 'timeZone': 'America/Caracas'}
                    # Obtener paciente para rellenar la descripción
                    cursor.execute("SELECT nombres, apellidos, cedula FROM pacientes WHERE id = ?", (row['paciente_id'],))
                    pac = cursor.fetchone()
                    g_event['description'] = f"Cédula: {pac['cedula'] if pac else ''}\nModalidad: {tipo_consulta}\nEstado: {estado_pago}"
                    service.events().update(calendarId='primary', eventId=google_event_id, body=g_event).execute()
                except Exception as ge:
                    print("Error al sincronizar cambio con Google Calendar:", ge)
        
        confirmada = data.get('confirmada') if 'confirmada' in data else row['confirmada']
        
        cursor.execute("""
            UPDATE agenda_finanzas SET
                estado_pago = ?,
                control_uso = ?,
                fecha_liquidacion = ?,
                monto = ?,
                moneda = ?,
                cantidad_sesiones = ?,
                referencia = ?,
                metodo_pago = ?,
                fecha_pago = ?,
                fecha = ?,
                hora = ?,
                tipo_consulta = ?,
                confirmada = ?
            WHERE id = ?
        """, (
            estado_pago,
            control_uso,
            fecha_liquidacion,
            monto,
            moneda,
            cantidad_sesiones,
            referencia,
            metodo_pago,
            fecha_pago,
            fecha,
            hora,
            tipo_consulta,
            confirmada,
            trans_id
        ))
        db.commit()
        
        # Sincronización en segundo plano con Firebase
        import threading
        threading.Thread(target=sync_patient_to_firebase, args=(row['paciente_id'],)).start()
        
        return jsonify({'success': 'Transacción actualizada con éxito.'})
    except Exception as e:
        return jsonify({'error': f'Error al actualizar transacción: {str(e)}'}), 500


# ==========================================
# AGENDA Y GOOGLE CALENDAR
# ==========================================



@admin_bp.route('/api/admin/clear-all-data', methods=['POST'])
@login_required
def clear_all_data():
    """Borra todos los datos del psicólogo previa confirmación explícita escribiendo CONFIRMAR."""
    data = request.json or {}
    confirmation = str(data.get('confirmation', '')).strip().upper()
    
    if confirmation != 'CONFIRMAR':
        return jsonify({'error': 'Debes escribir "CONFIRMAR" para autorizar esta acción.'}), 400

    db = get_db()
    cursor = db.cursor()
    psicologo_id = session.get('user_id')

    try:
        # Obtener IDs de pacientes del psicólogo
        cursor.execute("SELECT id FROM pacientes WHERE psicologo_id = ?", (psicologo_id,))
        patient_ids = [r[0] for r in cursor.fetchall()]

        if patient_ids:
            placeholders = ','.join('?' for _ in patient_ids)
            cursor.execute(f"DELETE FROM sesiones WHERE paciente_id IN ({placeholders})", patient_ids)
            cursor.execute(f"DELETE FROM agenda_finanzas WHERE paciente_id IN ({placeholders})", patient_ids)
            cursor.execute(f"DELETE FROM pacientes WHERE id IN ({placeholders})", patient_ids)

        cursor.execute("DELETE FROM pizarra_visual WHERE psicologo_id = ?", (psicologo_id,))
        cursor.execute("DELETE FROM notificaciones")

        db.commit()

        return jsonify({'success': 'Todos los datos de tu consultorio han sido eliminados con éxito.'})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al eliminar datos: {str(e)}'}), 500




@admin_bp.route('/api/firebase/config', methods=['GET'])
def get_firebase_config():
    _def_cfg = json.dumps({
        "apiKey": "AIzaSyDRQlUEv1SToy5ZdQQyUuYZDIhejeJ81zM",
        "authDomain": "espacio-terapeutico.firebaseapp.com",
        "databaseURL": "https://espacio-terapeutico-default-rtdb.firebaseio.com",
        "projectId": "espacio-terapeutico",
        "storageBucket": "espacio-terapeutico.firebasestorage.app",
        "messagingSenderId": "437385369836",
        "appId": "1:437385369836:web:f3745dc8d65d7ca418edc9",
        "measurementId": "G-M04FWL2963"
    })
    _def_vapid = "BIexDrYPs7iSYmxpkfgQwzatXm_o5pRa1ZAZUvzeF40nAc8N61RFlHqlZ153VNamBelgsKhB4nnowPJm_7Y-Qjc"

    cfg_val = _def_cfg
    vapid_val = _def_vapid

    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SELECT valor FROM configuracion WHERE clave = 'firebase_config'")
        row_cfg = cursor.fetchone()
        cursor.execute("SELECT valor FROM configuracion WHERE clave = 'firebase_vapid_key'")
        row_vapid = cursor.fetchone()
        if row_cfg and row_cfg[0]:
            cfg_val = row_cfg[0]
        if row_vapid and row_vapid[0]:
            vapid_val = row_vapid[0]
    except Exception as e:
        print("Error leyendo configuracion DB:", e)

    try:
        parsed_cfg = json.loads(cfg_val)
    except Exception:
        parsed_cfg = json.loads(_def_cfg)

    # Asegurar que el apiKey coincida con el apiKey oficial de Firebase Console
    parsed_cfg["apiKey"] = "AIzaSyDRQlUEv1SToy5ZdQQyUuYZDIhejeJ81zM"
    cfg_val = json.dumps(parsed_cfg)

    return jsonify({
        "config": cfg_val,
        "vapid_key": vapid_val,
        "vapidKey": vapid_val,
        "apiKey": parsed_cfg.get("apiKey", ""),
        "authDomain": parsed_cfg.get("authDomain", ""),
        "databaseURL": parsed_cfg.get("databaseURL", ""),
        "projectId": parsed_cfg.get("projectId", ""),
        "storageBucket": parsed_cfg.get("storageBucket", ""),
        "messagingSenderId": parsed_cfg.get("messagingSenderId", ""),
        "appId": parsed_cfg.get("appId", "")
    }), 200



@admin_bp.route('/api/firebase/config', methods=['POST'])
@login_required
def save_firebase_config():
    data = request.json or {}
    config_json = data.get('config')
    vapid_key = data.get('vapid_key')
    sa_json = data.get('sa_json')
    
    if config_json and vapid_key:
        try:
            import json
            json.loads(config_json) # Validar que sea JSON válido
            db = get_db()
            cursor = db.cursor()
            cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES ('firebase_config', ?)", (config_json,))
            cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES ('firebase_vapid_key', ?)", (vapid_key,))
            db.commit()
        except Exception as e:
            return jsonify({'error': f'Configuración SDK de Firebase Web no es un JSON válido: {str(e)}'}), 400

    if sa_json and sa_json.strip():
        try:
            import json
            config_data = json.loads(sa_json.strip())
            if 'private_key' in config_data and 'client_email' in config_data:
                with open(FIREBASE_SA_FILE, 'w', encoding='utf-8') as f:
                    json.dump(config_data, f, indent=4)
            else:
                return jsonify({'error': 'El JSON pegado no corresponde a una cuenta de servicio válida.'}), 400
        except Exception as e:
            return jsonify({'error': f'JSON de cuenta de servicio inválido: {str(e)}'}), 400

    return jsonify({'success': 'Configuración de Firebase guardada con éxito.'})



@admin_bp.route('/api/firebase/upload-sa', methods=['POST'])
@login_required
def upload_firebase_sa():
    if 'file' not in request.files:
        return jsonify({'error': 'No se proporcionó ningún archivo.'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío.'}), 400
    if not file.filename.endswith('.json'):
        return jsonify({'error': 'El archivo debe ser en formato JSON.'}), 400
    try:
        import json
        content = file.read().decode('utf-8')
        config_data = json.loads(content)
        # Validar estructura de cuenta de servicio de Firebase / Google Cloud
        if 'private_key' not in config_data or 'client_email' not in config_data:
            return jsonify({'error': 'El archivo no es una cuenta de servicio de Firebase válida.'}), 400
        
        with open(FIREBASE_SA_FILE, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, indent=4)
            
        return jsonify({'success': 'Cuenta de servicio de Firebase subida e instalada con éxito.'})
    except Exception as e:
        return jsonify({'error': f'Error al procesar el archivo: {str(e)}'}), 500



@admin_bp.route('/api/firebase/save-sa-text', methods=['POST'])
@login_required
def save_firebase_sa_text():
    data = request.json or {}
    sa_text = data.get('sa_json')
    if not sa_text:
        return jsonify({'error': 'El contenido del JSON es requerido.'}), 400
    try:
        import json
        config_data = json.loads(sa_text)
        if 'private_key' not in config_data or 'client_email' not in config_data:
            return jsonify({'error': 'El texto ingresado no corresponde a una cuenta de servicio de Firebase válida.'}), 400
            
        with open(FIREBASE_SA_FILE, 'w', encoding='utf-8') as f:
            json.dump(config_data, f, indent=4)
            
        return jsonify({'success': 'Cuenta de servicio de Firebase guardada con éxito.'})
    except Exception as e:
        return jsonify({'error': f'JSON inválido: {str(e)}'}), 500



@admin_bp.route('/api/firebase/status', methods=['GET'])
@login_required
def get_firebase_status():
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT valor FROM configuracion WHERE clave = 'firebase_config'")
    has_config = cursor.fetchone() is not None
    has_sa = os.path.exists(FIREBASE_SA_FILE)
    return jsonify({
        'configured': has_config,
        'has_service_account': has_sa
    })



@admin_bp.route('/api/firebase/subscribe', methods=['POST'])
def subscribe_firebase():
    data = request.json or {}
    token = data.get('token')
    if not token:
        return jsonify({'error': 'Token FCM requerido.'}), 400

    user_id = session.get('user_id')
    patient_id = session.get('patient_id')

    db = get_db()
    cursor = db.cursor()

    if user_id:
        # Actualizar cualquier token anónimo existente (NULL) o insertar con user_id
        cursor.execute("UPDATE fcm_subscriptions SET user_id = ? WHERE token = ?", (user_id, token))
        if cursor.rowcount == 0:
            cursor.execute("""
                INSERT OR REPLACE INTO fcm_subscriptions (user_id, patient_id, token)
                VALUES (?, NULL, ?)
            """, (user_id, token))
    elif patient_id:
        # Actualizar o insertar con patient_id
        cursor.execute("UPDATE fcm_subscriptions SET patient_id = ? WHERE token = ?", (patient_id, token))
        if cursor.rowcount == 0:
            cursor.execute("""
                INSERT OR REPLACE INTO fcm_subscriptions (user_id, patient_id, token)
                VALUES (NULL, ?, ?)
            """, (patient_id, token))
    else:
        # Sin sesión activa: guardar como anónimo (se actualizará al hacer login)
        cursor.execute("""
            INSERT OR REPLACE INTO fcm_subscriptions (user_id, patient_id, token)
            VALUES (NULL, NULL, ?)
        """, (token,))

    db.commit()
    return jsonify({'success': 'Suscrito a notificaciones FCM con éxito.'})



@admin_bp.route('/api/google/authorize')
@login_required
def google_authorize():
    import traceback
    try:
        if not GOOGLE_CALENDAR_AVAILABLE:
            return jsonify({
                'error': 'Fallo al iniciar flujo con Google Calendar',
                'detalle': 'Las librerías de Google Calendar no están instaladas en PythonAnywhere. Ejecuta pip install google-auth-oauthlib google-api-python-client en la consola.'
            }), 500

        if not os.path.exists(CLIENT_SECRETS_FILE):
            return "Error: Falta el archivo credentials.json en el servidor.", 400
            
        redirect_uri = url_for('admin.google_callback', _external=True)
        if not redirect_uri.startswith('https://') and 'localhost' not in redirect_uri and '127.0.0.1' not in redirect_uri:
            redirect_uri = redirect_uri.replace('http://', 'https://')
            
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            redirect_uri=redirect_uri
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true',
            prompt='consent'
        )
        
        session['state'] = state
        if getattr(flow, 'code_verifier', None):
            session['code_verifier'] = flow.code_verifier

        return redirect(authorization_url)
    except Exception as e:
        print("Error en google_authorize:", traceback.format_exc())
        return jsonify({
            'error': 'Fallo al iniciar flujo con Google Calendar',
            'detalle': str(e)
        }), 500



@admin_bp.route('/api/google/callback')
def google_callback():
    import traceback
    try:
        if not GOOGLE_CALENDAR_AVAILABLE:
            return "Error: Librerías de Google no instaladas.", 500

        state = session.get('state') or request.args.get('state')
        
        redirect_uri = url_for('admin.google_callback', _external=True)
        if not redirect_uri.startswith('https://') and 'localhost' not in redirect_uri and '127.0.0.1' not in redirect_uri:
            redirect_uri = redirect_uri.replace('http://', 'https://')
            
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            state=state,
            redirect_uri=redirect_uri
        )
        
        if 'code_verifier' in session and session['code_verifier']:
            flow.code_verifier = session['code_verifier']
        else:
            flow.code_verifier = None

        req_url = request.url
        if not req_url.startswith('https://') and 'localhost' not in req_url and '127.0.0.1' not in req_url:
            req_url = req_url.replace('http://', 'https://')
            
        os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
        os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'
        
        try:
            flow.fetch_token(authorization_response=req_url)
        except Exception as token_err:
            err_str = str(token_err)
            if 'mismatching_state' in err_str or 'State not equal' in err_str or 'Missing code verifier' in err_str or 'invalid_grant' in err_str:
                url_state = request.args.get('state')
                flow = Flow.from_client_secrets_file(
                    CLIENT_SECRETS_FILE,
                    scopes=SCOPES,
                    state=url_state,
                    redirect_uri=redirect_uri
                )
                flow.code_verifier = None
                auth_code = request.args.get('code')
                if auth_code:
                    flow.fetch_token(code=auth_code)
                else:
                    raise token_err
            else:
                raise token_err
                
        creds = flow.credentials
        
        # Guardar en base de datos local
        db = get_db()
        cursor = db.cursor()
        user_id = session.get('user_id')
        token_key = f'google_token_{user_id}' if user_id else 'google_token'
        cursor.execute("INSERT OR REPLACE INTO configuracion (clave, valor) VALUES (?, ?)", 
                       (token_key, creds.to_json()))
        db.commit()
        
        return """
        <html>
            <body style="font-family: sans-serif; text-align: center; padding: 2rem;">
                <h2 style="color: #059669;">✓ Conexión con Google Calendar exitosa</h2>
                <p>Tu cuenta de Google ha sido vinculada correctamente. Esta ventana se cerrará en breve.</p>
                <script>
                    if (window.opener) {
                        try { window.opener.location.reload(); } catch(e) {}
                    }
                    setTimeout(() => window.close(), 2000);
                </script>
            </body>
        </html>
        """
    except Exception as e:
        print("Error en google_callback:", traceback.format_exc())
        return f"""
        <html>
            <body style="font-family: sans-serif; text-align: center; padding: 2rem;">
                <h3 style="color: #dc2626;">Fallo al completar la autorización con Google Calendar</h3>
                <p>Detalle del error: {str(e)}</p>
                <button onclick="window.close()" style="padding: 0.5rem 1rem; border-radius: 6px; background: #374151; color: white; border: none; cursor: pointer;">Cerrar Ventana</button>
            </body>
        </html>
        """, 500



@admin_bp.route('/api/google/sync', methods=['POST'])
@login_required
def sync_google_calendar():
    import traceback
    try:
        user_id = session.get('user_id')
        service = get_calendar_service(user_id)
        if not service:
            return jsonify({'error': 'Google Calendar no está configurado o autorizado.'}), 400
            
        db = get_db()
        cursor = db.cursor()
        
        # 1. Traer eventos futuros de Google Calendar
        now = datetime.datetime.utcnow().isoformat() + 'Z' # 'Z' indica UTC
        events_result = service.events().list(
            calendarId='primary', timeMin=now,
            maxResults=100, singleEvents=True,
            orderBy='startTime'
        ).execute()
        g_events = events_result.get('items', [])
        
        synced_count = 0
        for ge in g_events:
            g_id = ge['id']
            summary = ge.get('summary', '')
            desc = ge.get('description', '')
            
            # Buscar si el evento ya está sincronizado localmente
            cursor.execute("SELECT id FROM agenda_finanzas WHERE google_event_id = ?", (g_id,))
            local_event = cursor.fetchone()
            
            if local_event:
                # Ya existe localmente. Actualizamos fecha/hora si cambió
                start = ge['start'].get('dateTime') or ge['start'].get('date')
                if start and 'T' in start:
                    fecha_g = start.split('T')[0]
                    hora_g = start.split('T')[1][:5]
                    
                    cursor.execute("""
                        UPDATE agenda_finanzas 
                        SET fecha = ?, hora = ? 
                        WHERE google_event_id = ?
                    """, (fecha_g, hora_g, g_id))
                synced_count += 1
            else:
                # Es nuevo desde Google Calendar. Intentamos enlazarlo a un paciente por nombre
                paciente_id = None
                if "Consulta:" in summary or "Consulta Psicológica -" in summary:
                    nombre_buscado = summary.replace("Consulta Psicológica -", "").replace("Consulta:", "").strip()
                    cursor.execute("""
                        SELECT id FROM pacientes 
                        WHERE (nombres || ' ' || apellidos) LIKE ? 
                        LIMIT 1
                    """, (f"%{nombre_buscado}%",))
                    pac_row = cursor.fetchone()
                    if pac_row:
                        paciente_id = pac_row['id']
                
                if paciente_id:
                    start = ge['start'].get('dateTime') or ge['start'].get('date')
                    if start and 'T' in start:
                        fecha_g = start.split('T')[0]
                        hora_g = start.split('T')[1][:5]
                        
                        modalidad = 'Online'
                        if 'Presencial' in desc or 'presencial' in summary.lower():
                            modalidad = 'Presencial'
                            
                        cursor.execute("""
                            INSERT INTO agenda_finanzas (
                                paciente_id, fecha, hora, tipo_consulta, monto, moneda, 
                                estado_pago, control_uso, google_event_id
                            ) VALUES (?, ?, ?, ?, 0.0, 'USD', 'Pendiente', 'Consumida', ?)
                        """, (paciente_id, fecha_g, hora_g, modalidad, g_id))
                        synced_count += 1

        # 2. Exportar citas futuras locales que no tengan google_event_id hacia Google Calendar
        today_str = datetime.datetime.now().strftime("%Y-%m-%d")
        cursor.execute("""
            SELECT af.id, af.fecha, af.hora, af.tipo_consulta, p.nombres, p.apellidos, p.email
            FROM agenda_finanzas af
            JOIN pacientes p ON af.paciente_id = p.id
            WHERE af.google_event_id IS NULL
              AND af.fecha >= ?
              AND (af.estado_pago IS NULL OR (af.estado_pago NOT LIKE 'Cancelada%' AND af.estado_pago != 'Reprogramada'))
              AND (p.psicologo_id = ? OR ? IS NULL)
        """, (today_str, user_id, user_id))
        local_pending = cursor.fetchall()
        
        pushed_count = 0
        for lp in local_pending:
            try:
                pac_nombre = f"{lp['nombres']} {lp['apellidos']}"
                start_dt = f"{lp['fecha']}T{lp['hora']}:00-04:00"
                end_h = str(int(lp['hora'].split(':')[0]) + 1).zfill(2)
                end_dt = f"{lp['fecha']}T{end_h}:{lp['hora'].split(':')[1]}:00-04:00"
                
                event_b = {
                    'summary': f"Consulta Psicológica - {pac_nombre}",
                    'description': f"Modalidad: {lp['tipo_consulta']}",
                    'start': {'dateTime': start_dt, 'timeZone': 'America/Caracas'},
                    'end': {'dateTime': end_dt, 'timeZone': 'America/Caracas'}
                }
                if lp['email']:
                    event_b['attendees'] = [{'email': lp['email'], 'displayName': pac_nombre}]
                    
                g_ev = service.events().insert(calendarId='primary', body=event_b, sendUpdates='all').execute()
                if g_ev.get('id'):
                    cursor.execute("UPDATE agenda_finanzas SET google_event_id = ? WHERE id = ?", (g_ev['id'], lp['id']))
                    pushed_count += 1
            except Exception as pe:
                print("Error enviando cita local a Google Calendar:", pe)

        db.commit()
        
        msg = f"✓ Sincronización completada exitosamente."
        if pushed_count > 0:
            msg += f" {pushed_count} cita(s) enviada(s) a Google Calendar."
        if synced_count > 0:
            msg += f" {synced_count} evento(s) actualizado(s)/importado(s)."
        if pushed_count == 0 and synced_count == 0:
            msg += " Tu agenda ya estaba 100% al día."
            
        return jsonify({'success': msg})
        
    except Exception as e:
        print("Error durante la sincronización de Google Calendar:", traceback.format_exc())
        return jsonify({'error': f'Error durante la sincronización: {str(e)}'}), 500


# ==========================================
# EXPORTACIÓN A WORD (.DOCX)
# ==========================================



@admin_bp.route('/api/export/word/<int:patient_id>', methods=['GET'])
@login_required
def export_word(patient_id):
    db = get_db()
    cursor = db.cursor()
    
    # 1. Obtener datos del paciente
    cursor.execute("SELECT * FROM pacientes WHERE id = ?", (patient_id,))
    pac = cursor.fetchone()
    if not pac:
        return jsonify({'error': 'Paciente no encontrado'}), 404
        
    # 2. Obtener sesiones
    cursor.execute("SELECT * FROM sesiones WHERE paciente_id = ? ORDER BY fecha ASC", (patient_id,))
    sessions = cursor.fetchall()
    
    # 3. Obtener balance financiero
    cursor.execute("SELECT * FROM agenda_finanzas WHERE paciente_id = ? ORDER BY fecha ASC", (patient_id,))
    finance_events = cursor.fetchall()
    
    # Decodificar datos cifrados del paciente
    pac_dict = dict(pac)
    for field in ['antecedentes_medicos_personales', 'antecedentes_medicos_familiares', 
                  'antecedentes_psicologicos_personales', 'antecedentes_psicologicos_familiares',
                  'asistencia_previa_psicologo', 'motivo_consulta', 'expectativas', 
                  'farmacologia', 'diagnostico']:
        if pac_dict.get(field):
            pac_dict[field] = decrypt_clinical_text(pac_dict[field])
            
    # Crear documento Word
    doc = docx.Document()
    
    # Estilo y Título Principal
    title = doc.add_paragraph()
    title_run = title.add_run("HISTORIA CLÍNICA Y EXPEDIENTE PSICOLÓGICO")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = docx.shared.RGBColor(0x3D, 0x1E, 0x3F) # Berenjena Oscuro
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"Consultante: {pac_dict['nombres']} {pac_dict['apellidos']} | Cédula: {pac_dict['cedula']}")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Margen horizontal
    doc.add_paragraph("__________________________________________________________________")
    
    # Sección 1: Datos Personales
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Datos Personales y Filiación")
    h1_run.font.color.rgb = docx.shared.RGBColor(0x3D, 0x1E, 0x3F)
    
    table_data = [
        ("Nombres y Apellidos", f"{pac_dict['nombres']} {pac_dict['apellidos']}"),
        ("Cédula de Identidad", pac_dict['cedula']),
        ("Pronombre / Género", f"{pac_dict['pronombre'] or 'N/A'} / {pac_dict['genero'] or 'N/A'}"),
        ("Edad", str(pac_dict['edad']) if pac_dict['edad'] else "N/A"),
        ("Lugar y Fecha de Nacimiento", f"{pac_dict['lugar_nacimiento'] or 'N/A'} ({pac_dict['fecha_nacimiento'] or 'N/A'})"),
        ("Residencia Actual", ", ".join(filter(None, [pac_dict['residencia_actual'] if 'residencia_actual' in pac_dict.keys() else (pac_dict['ciudad'] if 'ciudad' in pac_dict.keys() else None), pac_dict['pais'] if 'pais' in pac_dict.keys() else None])) or "N/A"),
        ("Reside con", pac_dict['con_quien_reside'] or "N/A"),
        ("Nivel Académico / Ocupación", f"{pac_dict['nivel_academico'] or 'N/A'} / {pac_dict['ocupacion'] or 'N/A'}"),
        ("Estado Civil / Relacional", pac_dict['estado_civil'] or "N/A"),
    ]
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Shading Accent 1'
    for label, val in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = val
        
    doc.add_paragraph() # Espacio
    
    # Sección 2: Antecedentes e Impresión Diagnóstica
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Antecedentes e Impresión Diagnóstica")
    h2_run.font.color.rgb = docx.shared.RGBColor(0x3D, 0x1E, 0x3F)
    
    doc.add_paragraph().add_run("Antecedentes Médicos:").bold = True
    doc.add_paragraph(f"- Personales: {pac_dict['antecedentes_medicos_personales'] or 'Sin registrar'}")
    doc.add_paragraph(f"- Familiares: {pac_dict['antecedentes_medicos_familiares'] or 'Sin registrar'}")
    
    doc.add_paragraph().add_run("Antecedentes Psicológicos y Psiquiátricos:").bold = True
    doc.add_paragraph(f"- Personales: {pac_dict['antecedentes_psicologicos_personales'] or 'Sin registrar'}")
    doc.add_paragraph(f"- Familiares: {pac_dict['antecedentes_psicologicos_familiares'] or 'Sin registrar'}")
    
    doc.add_paragraph().add_run("Motivo de Consulta y Expectativas:").bold = True
    doc.add_paragraph(f"- Asistencia previa al psicólogo: {pac_dict['asistencia_previa_psicologo'] or 'Sin registrar'}")
    doc.add_paragraph(f"- Motivo de consulta actual: {pac_dict['motivo_consulta'] or 'Sin registrar'}")
    doc.add_paragraph(f"- Expectativas del proceso: {pac_dict['expectativas'] or 'Sin registrar'}")
    
    doc.add_paragraph().add_run("Tratamiento Farmacológico Activo:").bold = True
    doc.add_paragraph(pac_dict['farmacologia'] or "Ninguno")
    
    doc.add_paragraph().add_run("Contacto de Emergencia:").bold = True
    doc.add_paragraph(f"{pac_dict['contacto_emergencia_nombre'] or 'N/A'} ({pac_dict['contacto_emergencia_parentesco'] or 'N/A'})")
    
    doc.add_paragraph().add_run("Impresión Diagnóstica Evolutiva:").bold = True
    doc.add_paragraph(pac_dict['diagnostico'] or "Sin impresión diagnóstica anotada.")
    
    doc.add_page_break()
    
    # Sección 3: Evolución Cronológica (Sesiones)
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Registro de Sesiones (Evolución)")
    h3_run.font.color.rgb = docx.shared.RGBColor(0x3D, 0x1E, 0x3F)
    
    if not sessions:
        doc.add_paragraph("No hay sesiones de evolución registradas para este consultante.")
    else:
        for idx, s in enumerate(sessions, 1):
            s_resumen = decrypt_clinical_text(s['resumen']) or "Sin resumen."
            s_tareas = decrypt_clinical_text(s['tareas_asignadas']) or "Ninguna."
            s_recursos = decrypt_clinical_text(s['recursos_entregados']) or "Ninguno."
            s_anotaciones = decrypt_clinical_text(s['anotaciones_proxima']) or "Ninguna."
            s_compromisos = decrypt_clinical_text(s['compromisos_psicologo']) or "Ninguno."
            
            p_ses = doc.add_paragraph()
            p_ses.add_run(f"Sesión N° {idx} — Fecha: {s['fecha']} | Modalidad: {s['modalidad']}").bold = True
            doc.add_paragraph().add_run("Resumen abordado:").bold = True
            doc.add_paragraph(s_resumen)
            doc.add_paragraph().add_run("Tareas asignadas al consultante:").bold = True
            doc.add_paragraph(s_tareas)
            doc.add_paragraph().add_run("Recursos entregados:").bold = True
            doc.add_paragraph(s_recursos)
            doc.add_paragraph().add_run("Anotaciones próxima consulta:").bold = True
            doc.add_paragraph(s_anotaciones)
            doc.add_paragraph().add_run("Compromisos del psicólogo:").bold = True
            doc.add_paragraph(s_compromisos)
            doc.add_paragraph("____________________________________________________")
            
    doc.add_page_break()
    
    # Sección 4: Historial de Citas y Finanzas
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Historial de Citas y Estado de Cuentas")
    h4_run.font.color.rgb = docx.shared.RGBColor(0x3D, 0x1E, 0x3F)
    
    if not finance_events:
        doc.add_paragraph("No hay registro de citas o transacciones financieras asociadas.")
    else:
        table_f = doc.add_table(rows=1, cols=6)
        table_f.style = 'Light Shading Accent 1'
        hdr_cells = table_f.rows[0].cells
        hdr_cells[0].text = 'Fecha'
        hdr_cells[1].text = 'Hora'
        hdr_cells[2].text = 'Modalidad'
        hdr_cells[3].text = 'Monto'
        hdr_cells[4].text = 'Estado Pago'
        hdr_cells[5].text = 'Control Uso'
        
        for fe in finance_events:
            row_cells = table_f.add_row().cells
            row_cells[0].text = fe['fecha']
            row_cells[1].text = fe['hora']
            row_cells[2].text = fe['tipo_consulta']
            row_cells[3].text = f"{fe['monto']} {fe['moneda']}"
            row_cells[4].text = fe['estado_pago']
            row_cells[5].text = fe['control_uso']
            
    # Guardar en archivo temporal
    filename = f"expediente_{pac['cedula']}.docx"
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    
    # Enviar archivo
    return send_file(filepath, as_attachment=True, download_name=filename)


# ==========================================
# COPIAS DE SEGURIDAD / RESPALDO
# ==========================================



@admin_bp.route('/api/backup', methods=['GET'])
@login_required
def create_backup():
    """Descarga la base de datos .db directamente al navegador."""
    if not os.path.exists(DATABASE):
        return jsonify({'error': 'La base de datos aún no se ha inicializado.'}), 400
        
    dt = datetime.datetime.now()
    now_str = dt.strftime("%Y-%m-%d_%H-%M")
    backup_filename = f"copia_seguridad_clinica_{now_str}.db"
    
    try:
        db = getattr(g, '_database', None)
        if db is not None:
            db.close()
            
        return send_file(
            DATABASE,
            as_attachment=True,
            download_name=backup_filename,
            mimetype='application/x-sqlite3'
        )
    except Exception as e:
        return jsonify({'error': f'Error al descargar copia de seguridad: {str(e)}'}), 500



@admin_bp.route('/api/admin/backup/export-patients-word-zip', methods=['GET'])
@login_required
def export_patients_word_zip():
    db = get_db()
    cursor = db.cursor()
    user_id = session.get('user_id', 1)
    
    cursor.execute("SELECT * FROM pacientes WHERE psicologo_id = ?", (user_id,))
    patients = cursor.fetchall()
    
    if not patients:
        return jsonify({'error': 'No hay pacientes para exportar.'}), 404
        
    import io, zipfile, datetime
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

    def set_cell_bg(cell, hex_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for p_row in patients:
            p = dict(p_row)
            doc = Document()

            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_title = title_p.add_run(f"EXPEDIENTE CLÍNICO DE {p.get('nombres','').upper()} {p.get('apellidos','').upper()}")
            r_title.bold = True
            r_title.font.size = Pt(16)
            r_title.font.color.rgb = RGBColor(16, 185, 129)

            h1 = doc.add_paragraph()
            r_h1 = h1.add_run("1. DATOS DE IDENTIFICACIÓN Y ANTECEDENTES")
            r_h1.bold = True
            r_h1.font.size = Pt(12)

            table = doc.add_table(rows=0, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            fields = [
                ("Nombre Completo", f"{p.get('nombres','')} {p.get('apellidos','')}"),
                ("Cédula / ID", p.get('cedula','N/A')),
                ("Teléfono", p.get('telefono','N/A')),
                ("Correo Electrónico", p.get('email','N/A')),
                ("Edad / Género", f"{p.get('edad','N/A')} años / {p.get('genero','N/A')}"),
                ("Estado Civil", p.get('estado_civil','N/A')),
                ("Ocupación", p.get('ocupacion','N/A')),
                ("Motivo de Consulta", p.get('motivo_consulta','N/A')),
                ("Diagnóstico Inicial", p.get('diagnostico','Sin diagnóstico registrado'))
            ]
            for label, val in fields:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(val or 'N/A')
                set_cell_bg(row_cells[0], "F8FAFC")

            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            raw_filename = f"Expediente_{p.get('cedula', 'ID')}_{p.get('nombres', '')}_{p.get('apellidos', '')}.docx"
            safe_filename = re.sub(r'[\\/*?:"<>|]', '_', raw_filename).replace(' ', '_')
            
            doc_stream = io.BytesIO()
            doc.save(doc_stream)
            zip_file.writestr(safe_filename, doc_stream.getvalue())

    zip_buffer.seek(0)
    today_str = datetime.date.today().strftime('%Y-%m-%d')
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'Expedientes_Pacientes_EspacioTerapeutico_{today_str}.zip'
    )

@admin_bp.route('/api/restore', methods=['POST'])
@login_required
def restore_backup():
    if 'file' not in request.files:
        return jsonify({'error': 'No se cargó ningún archivo.'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío.'}), 400
        
    if not file.filename.endswith('.db'):
        return jsonify({'error': 'El archivo de respaldo debe tener extensión .db'}), 400

    import tempfile, sqlite3 as _sqlite3

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.db')
    file.save(tmp.name)
    tmp.close()

    stats = {'pacientes': 0, 'agenda': 0, 'sesiones': 0, 'omitidos': 0, 'errores': []}

    try:
        conn_b = _sqlite3.connect(tmp.name)
        conn_b.row_factory = _sqlite3.Row
        cur_b = conn_b.cursor()

        cur_b.execute("SELECT name FROM sqlite_master WHERE type='table'")
        backup_tables = {r[0] for r in cur_b.fetchall()}

        db_target = get_db()
        cur_t = db_target.cursor()
        cur_t.execute("SELECT id FROM usuarios WHERE id = ?", (session.get('user_id', 1),))
        psic_row = cur_t.fetchone()
        psic_id = psic_row['id'] if psic_row else 1

        def backup_cols(table_name):
            try:
                cur_b.execute(f"PRAGMA table_info(`{table_name}`)")
                return {r['name'] for r in cur_b.fetchall()}
            except Exception:
                return set()

        if 'pacientes' in backup_tables:
            cols_b = backup_cols('pacientes')
            cur_b.execute("SELECT * FROM pacientes")
            for p in cur_b.fetchall():
                p = dict(p)
                cedula = p.get('cedula') or ''
                cur_t.execute(
                    "SELECT id FROM pacientes WHERE id=? OR (cedula!='' AND cedula=?)",
                    (p['id'], cedula)
                )
                if cur_t.fetchone():
                    stats['omitidos'] += 1
                    continue
                try:
                    cur_t.execute("""
                        INSERT INTO pacientes (
                            id, nombres, apellidos, cedula, pronombre, genero, edad,
                            lugar_nacimiento, fecha_nacimiento, residencia_actual,
                            con_quien_reside, nivel_academico, ocupacion, estado_civil,
                            antecedentes_medicos_familiares, antecedentes_medicos_personales,
                            antecedentes_psicologicos_familiares, antecedentes_psicologicos_personales,
                            asistencia_previa_psicologo, motivo_consulta, expectativas,
                            farmacologia, contacto_emergencia_nombre, contacto_emergencia_parentesco,
                            diagnostico, fecha_registro, telefono, email, psicologo_id
                        ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    """, (
                        p['id'],
                        p.get('nombres', ''),
                        p.get('apellidos', ''),
                        p.get('cedula', ''),
                        p.get('pronombre', ''),
                        p.get('genero', ''),
                        p.get('edad', ''),
                        p.get('lugar_nacimiento', ''),
                        p.get('fecha_nacimiento', ''),
                        p.get('residencia_actual', ''),
                        p.get('con_quien_reside', ''),
                        p.get('nivel_academico', ''),
                        p.get('ocupacion', ''),
                        p.get('estado_civil', ''),
                        p.get('antecedentes_medicos_familiares', ''),
                        p.get('antecedentes_medicos_personales', ''),
                        p.get('antecedentes_psicologicos_familiares', ''),
                        p.get('antecedentes_psicologicos_personales', ''),
                        p.get('asistencia_previa_psicologo', ''),
                        p.get('motivo_consulta', ''),
                        p.get('expectativas', ''),
                        p.get('farmacologia', ''),
                        p.get('contacto_emergencia_nombre', ''),
                        p.get('contacto_emergencia_parentesco', ''),
                        p.get('diagnostico', ''),
                        p.get('fecha_registro', ''),
                        p.get('telefono', ''),
                        p.get('email', ''),
                        p.get('psicologo_id', psic_id)
                    ))
                    stats['pacientes'] += 1
                except Exception as e:
                    stats['errores'].append(f"Paciente {p.get('nombres','?')}: {str(e)[:60]}")

        if 'agenda_finanzas' in backup_tables:
            cur_b.execute("SELECT * FROM agenda_finanzas")
            for a in cur_b.fetchall():
                a = dict(a)
                cur_t.execute("SELECT id FROM agenda_finanzas WHERE id=?", (a['id'],))
                if cur_t.fetchone():
                    stats['omitidos'] += 1
                    continue
                try:
                    cur_t.execute("""
                        INSERT INTO agenda_finanzas (
                            id, paciente_id, fecha, hora, google_event_id, tipo_consulta,
                            monto, moneda, estado_pago, control_uso, fecha_liquidacion,
                            cantidad_sesiones, referencia, metodo_pago, fecha_pago, confirmada
                        ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    """, (
                        a['id'], a.get('paciente_id'), a.get('fecha',''),
                        a.get('hora',''), a.get('google_event_id',''),
                        a.get('tipo_consulta','Individual'),
                        a.get('monto', 0), a.get('moneda','USD'),
                        a.get('estado_pago','Pendiente'),
                        a.get('control_uso', 0), a.get('fecha_liquidacion',''),
                        a.get('cantidad_sesiones', 1), a.get('referencia',''),
                        a.get('metodo_pago',''), a.get('fecha_pago',''),
                        a.get('confirmada', 1)
                    ))
                    stats['agenda'] += 1
                except Exception as e:
                    stats['errores'].append(f"Agenda id={a.get('id')}: {str(e)[:60]}")

        if 'sesiones' in backup_tables:
            from app import encrypt_clinical_text
            cur_b.execute("SELECT * FROM sesiones")
            for s in cur_b.fetchall():
                s = dict(s)
                cur_t.execute("SELECT id FROM sesiones WHERE id=?", (s['id'],))
                if cur_t.fetchone():
                    stats['omitidos'] += 1
                    continue

                def _safe_enc(val):
                    if not val:
                        return ''
                    v = str(val)
                    if v.startswith('enc:'):
                        return v
                    return encrypt_clinical_text(v)

                try:
                    cur_t.execute("""
                        INSERT INTO sesiones (
                            id, paciente_id, agenda_id, fecha, modalidad, estado,
                            resumen, tareas_asignadas, recursos_entregados,
                            anotaciones_proxima, compromisos_psicologo,
                            diagnostico, test_aplicados, archivo_adjunto
                        ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    """, (
                        s['id'], s.get('paciente_id'), s.get('agenda_id'),
                        s.get('fecha',''), s.get('modalidad','Online'),
                        s.get('estado','Realizada'),
                        _safe_enc(s.get('resumen')),
                        s.get('tareas_asignadas',''),
                        s.get('recursos_entregados',''),
                        _safe_enc(s.get('anotaciones_proxima')),
                        _safe_enc(s.get('compromisos_psicologo')),
                        _safe_enc(s.get('diagnostico')),
                        _safe_enc(s.get('test_aplicados')),
                        s.get('archivo_adjunto','')
                    ))
                    stats['sesiones'] += 1
                except Exception as e:
                    stats['errores'].append(f"Sesión id={s.get('id')}: {str(e)[:60]}")

        db_target.commit()
        conn_b.close()
        try:
            os.unlink(tmp.name)
        except Exception:
            pass

        msg = (f"Restauración completada: "
               f"{stats['pacientes']} pacientes, "
               f"{stats['agenda']} registros financieros, "
               f"{stats['sesiones']} sesiones importadas. "
               f"{stats['omitidos']} registros ya existían (omitidos).")
        if stats['errores']:
            msg += f" Advertencias: {'; '.join(stats['errores'][:3])}"

        return jsonify({'success': msg, 'stats': stats})

    except Exception as e:
        try:
            os.unlink(tmp.name)
        except Exception:
            pass
        return jsonify({'error': f'Error al restaurar: {str(e)}'}), 500

@admin_bp.after_request
def after_request(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return response


@admin_bp.route('/api/onboarding/skip', methods=['POST'])
@login_required
def skip_onboarding():
    db = get_db()
    cursor = db.cursor()
    user_id = session.get('user_id')
    cursor.execute("UPDATE usuarios SET primer_inicio = 0 WHERE id = ?", (user_id,))
    return jsonify({'success': 'Onboarding omitido con éxito.'})

@admin_bp.route('/api/onboarding/complete', methods=['POST'])
@login_required
def complete_onboarding():
    data = request.json or {}
    user_id = session.get('user_id')
    if not user_id:
        return jsonify({'error': 'Sesión no válida.'}), 401

    db = get_db()
    cursor = db.cursor()
    nombres = data.get('nombres')
    apellidos = data.get('apellidos')
    estudios = data.get('estudios', '')
    federacion = data.get('federacion', '')
    raw_slug = data.get('slug') or f"psic.{nombres}{apellidos}"
    cleaned_slug = raw_slug.strip().lower().replace(" ", "").replace("/", "").replace(".", "")
    if not cleaned_slug.startswith("psic"):
        cleaned_slug = "psic." + cleaned_slug
    else:
        cleaned_slug = "psic." + cleaned_slug[4:]
        
    duracion = int(data.get('duracion', 60))
    receso = int(data.get('receso', 15))
    perfiles = data.get('perfiles', [])
    metodos_pago = data.get('metodos_pago', {})
    
    if not nombres or not apellidos:
        return jsonify({'error': 'Nombres y Apellidos son obligatorios.'}), 400
        
    cursor.execute("SELECT id FROM usuarios WHERE slug = ? AND id != ?", (cleaned_slug, user_id))
    if cursor.fetchone():
        cleaned_slug = f"{cleaned_slug}{user_id}"

    default_visual = {
        "duracion": duracion,
        "receso": receso,
        "antelacion": 24,
        "alerta_confirmacion": 24,
        "alerta_recordatorio": 2,
        "alerta_cierre": 2,
        "limite_cancelacion_tipo": "horas",
        "limite_cancelacion_valor": 24,
        "perfiles": perfiles if perfiles else [
            {
                "id": "default_online",
                "nombre": "Horario Online",
                "modalidad": "Online",
                "dias": [
                    {"dia": 1, "nombre": "Lunes", "activo": True, "rangos": [{"inicio": "12:00", "fin": "16:00"}, {"inicio": "18:00", "fin": "22:00"}]},
                    {"dia": 2, "nombre": "Martes", "activo": True, "rangos": [{"inicio": "18:00", "fin": "22:00"}]},
                    {"dia": 3, "nombre": "Miércoles", "activo": False, "rangos": []},
                    {"dia": 4, "nombre": "Jueves", "activo": False, "rangos": []},
                    {"dia": 5, "nombre": "Viernes", "activo": False, "rangos": []},
                    {"dia": 6, "nombre": "Sábado", "activo": False, "rangos": []},
                    {"dia": 0, "nombre": "Domingo", "activo": False, "rangos": []}
                ]
            },
            {
                "id": "default_presencial",
                "nombre": "Horario Presencial",
                "modalidad": "Presencial",
                "dias": [
                    {"dia": 1, "nombre": "Lunes", "activo": False, "rangos": []},
                    {"dia": 2, "nombre": "Martes", "activo": False, "rangos": []},
                    {"dia": 3, "nombre": "Miércoles", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}]},
                    {"dia": 4, "nombre": "Jueves", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}]},
                    {"dia": 5, "nombre": "Viernes", "activo": True, "rangos": [{"inicio": "08:00", "fin": "12:00"}]},
                    {"dia": 6, "nombre": "Sábado", "activo": False, "rangos": []},
                    {"dia": 0, "nombre": "Domingo", "activo": False, "rangos": []}
                ]
            }
        ]
    }
    
    cfg_visual_str = json.dumps(default_visual)
    metodos_pago_str = json.dumps(metodos_pago) if metodos_pago else json.dumps({})

    try:
        cursor.execute("""
            UPDATE usuarios
            SET nombres = ?, apellidos = ?, estudios = ?, federacion = ?,
                slug = ?, configuracion_horarios_visual = ?, metodos_pago = ?,
                primer_inicio = 0
            WHERE id = ?
        """, (nombres, apellidos, estudios, federacion, cleaned_slug, cfg_visual_str, metodos_pago_str, user_id))
        db.commit()
        return jsonify({'success': '¡Bienvenido a tu consultorio! Configuración inicial completada.', 'slug': cleaned_slug})
    except Exception as e:
        return jsonify({'error': f'Error al guardar configuración inicial: {str(e)}'}), 500



@admin_bp.route('/api/superadmin/therapists/<int:user_id>', methods=['DELETE'])
@login_required
def superadmin_delete_therapist(user_id):
    if not check_is_superadmin():
        return jsonify({'error': 'Acceso denegado. Se requieren permisos de superadministrador.'}), 403
        
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT id, username, nombres, apellidos FROM usuarios WHERE id = ? AND role = 'psicologo'", (user_id,))
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Psicólogo no encontrado.'}), 404
        
    try:
        # Obtener pacientes del psicólogo antes de eliminar para limpiar Firebase
        cursor.execute("SELECT id, username, cedula FROM pacientes WHERE psicologo_id = ?", (user_id,))
        pats_to_delete = cursor.fetchall()
        
        # Cascade cleanup of all psychologist data
        cursor.execute("DELETE FROM pizarra_terapeutica WHERE paciente_id IN (SELECT id FROM pacientes WHERE psicologo_id = ?)", (user_id,))
        cursor.execute("DELETE FROM agenda_finanzas WHERE paciente_id IN (SELECT id FROM pacientes WHERE psicologo_id = ?)", (user_id,))
        cursor.execute("DELETE FROM sesiones WHERE paciente_id IN (SELECT id FROM pacientes WHERE psicologo_id = ?)", (user_id,))
        cursor.execute("DELETE FROM pacientes WHERE psicologo_id = ?", (user_id,))
        cursor.execute("DELETE FROM notificaciones WHERE user_id = ?", (user_id,))
        cursor.execute("DELETE FROM fcm_subscriptions WHERE user_id = ?", (user_id,))
        cursor.execute("DELETE FROM usuarios WHERE id = ?", (user_id,))
        db.commit()
        
        # Limpiar pacientes de Firebase Realtime Database
        for pt in pats_to_delete:
            delete_patient_from_firebase(pt['id'], (pt['username'] or pt['cedula'] or '').strip())
            
        return jsonify({'success': f"Psicólogo '{row['nombres']} {row['apellidos']}' (@{row['username']}) y toda su información fueron eliminados con éxito."})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al eliminar psicólogo: {str(e)}'}), 500



