# -*- coding: utf-8 -*-
"""
Módulo de Examen Mental Estructurado (MSE) (routes_examen_mental.py)
Encapsula la evaluación psicopatológica estandarizada (Apariencia, Porte, Actitud, Conciencia, Orientación,
Atención, Memoria, Pensamiento, Sensopercepción, Afecto, Juicio e Introspección) y la generación e impresión de informes en Word y PDF.
"""

import os
import io
import json
import sqlite3
from datetime import datetime
from functools import wraps
from flask import Blueprint, request, jsonify, session, g, Response, render_template_string

examen_mental_bp = Blueprint('examen_mental', __name__)

def get_db():
    """Obtiene la conexión a la base de datos desde el contexto global g de Flask."""
    db = getattr(g, '_database', None)
    if db is None:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(base_dir, 'clinica.db')
        db = g._database = sqlite3.connect(db_path, timeout=30.0)
        db.row_factory = sqlite3.Row
    return db

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({'error': 'No autorizado. Por favor inicia sesión.'}), 401
        return f(*args, **kwargs)
    return decorated_function

def _ensure_examenes_mentales_table(cursor):
    try:
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS examenes_mentales (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                psicologo_id INTEGER NOT NULL,
                paciente_id INTEGER NOT NULL,
                fecha_evaluacion TEXT NOT NULL,
                medio_evaluacion TEXT NOT NULL,
                datos_evaluacion_json TEXT NOT NULL,
                observaciones_generales TEXT,
                fecha_registro DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (psicologo_id) REFERENCES usuarios(id) ON DELETE CASCADE,
                FOREIGN KEY (paciente_id) REFERENCES pacientes(id) ON DELETE CASCADE
            )
        """)
    except Exception as _e:
        print("Error en _ensure_examenes_mentales_table:", _e)

@examen_mental_bp.route('/api/examen-mental', methods=['POST'])
@login_required
def save_examen_mental():
    user_id = session.get('user_id')
    db = get_db()
    cursor = db.cursor()
    
    if session.get('role') != 'superadmin':
        cursor.execute("SELECT COALESCE(bloqueo_examen_mental, 1) FROM usuarios WHERE id = ?", (user_id,))
        b_row = cursor.fetchone()
        if b_row and b_row[0] == 1:
            return jsonify({'error': 'La función de Examen Mental está inhabilitada para tu cuenta. Contacta a administración.'}), 403
    _ensure_examenes_mentales_table(cursor)
    
    data = request.json or {}
    paciente_id = data.get('paciente_id')
    fecha_evaluacion = data.get('fecha_evaluacion')
    medio_evaluacion = data.get('medio_evaluacion', 'Presencial')
    datos_evaluacion = data.get('datos_evaluacion_json', {})
    observaciones_generales = data.get('observaciones_generales', '').strip()
    
    if not paciente_id or not fecha_evaluacion:
        return jsonify({'error': 'El paciente y la fecha de evaluación son requeridos.'}), 400
        
    try:
        datos_json_str = json.dumps(datos_evaluacion, ensure_ascii=False)
        
        cursor.execute("""
            INSERT INTO examenes_mentales (psicologo_id, paciente_id, fecha_evaluacion, medio_evaluacion, datos_evaluacion_json, observaciones_generales)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (user_id, paciente_id, fecha_evaluacion, medio_evaluacion, datos_json_str, observaciones_generales))
        
        exam_id = cursor.lastrowid
        db.commit()
        
        return jsonify({'success': 'Examen mental guardado con éxito e integrado a la historia clínica.', 'id': exam_id})
    except Exception as e:
        db.rollback()
        return jsonify({'error': f'Error al guardar el examen mental: {str(e)}'}), 500

@examen_mental_bp.route('/api/examen-mental/historial', methods=['GET'])
@login_required
def get_examen_mental_historial():
    user_id = session.get('user_id')
    search = request.args.get('search', '').strip()
    paciente_id_param = request.args.get('paciente_id')
    
    db = get_db()
    cursor = db.cursor()
    _ensure_examenes_mentales_table(cursor)
    
    query = """
        SELECT e.id, e.psicologo_id, e.paciente_id, e.fecha_evaluacion, e.medio_evaluacion,
               e.datos_evaluacion_json, e.observaciones_generales, e.fecha_registro,
               p.nombres as pac_nombres, p.apellidos as pac_apellidos, p.cedula as pac_cedula,
               p.genero as pac_genero, p.fecha_nacimiento as pac_fecha_nacimiento
        FROM examenes_mentales e
        JOIN pacientes p ON e.paciente_id = p.id
        WHERE e.psicologo_id = ?
    """
    params = [user_id]
    
    if paciente_id_param:
        query += " AND e.paciente_id = ?"
        params.append(paciente_id_param)
        
    if search:
        query += " AND (p.nombres LIKE ? OR p.apellidos LIKE ? OR p.cedula LIKE ?)"
        s_term = f"%{search}%"
        params.extend([s_term, s_term, s_term])
        
    query += " ORDER BY e.fecha_registro DESC, e.id DESC"
    
    cursor.execute(query, params)
    rows = cursor.fetchall()
    
    result = []
    for r in rows:
        r_dict = dict(r)
        try:
            r_dict['datos_evaluacion'] = json.loads(r_dict['datos_evaluacion_json']) if r_dict['datos_evaluacion_json'] else {}
        except:
            r_dict['datos_evaluacion'] = {}
        result.append(r_dict)
        
    return jsonify(result)

@examen_mental_bp.route('/api/examen-mental/<int:exam_id>', methods=['GET'])
@login_required
def get_examen_mental_detail(exam_id):
    user_id = session.get('user_id')
    db = get_db()
    cursor = db.cursor()
    _ensure_examenes_mentales_table(cursor)
    
    cursor.execute("""
        SELECT e.*, p.nombres as pac_nombres, p.apellidos as pac_apellidos, p.cedula as pac_cedula,
               p.genero as pac_genero, p.fecha_nacimiento as pac_fecha_nacimiento, p.ocupacion as pac_ocupacion,
               u.nombres as psic_nombres, u.apellidos as psic_apellidos, u.estudios as psic_estudios, u.federacion as psic_federacion
        FROM examenes_mentales e
        JOIN pacientes p ON e.paciente_id = p.id
        JOIN usuarios u ON e.psicologo_id = u.id
        WHERE e.id = ? AND (e.psicologo_id = ? OR ? = 1)
    """, (exam_id, user_id, 1 if session.get('role') in ['admin', 'superadmin'] else 0))
    
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Examen mental no encontrado o sin permisos.'}), 404
        
    r_dict = dict(row)
    try:
        r_dict['datos_evaluacion'] = json.loads(r_dict['datos_evaluacion_json']) if r_dict['datos_evaluacion_json'] else {}
    except:
        r_dict['datos_evaluacion'] = {}
        
    return jsonify(r_dict)

@examen_mental_bp.route('/api/examen-mental/<int:exam_id>/export/pdf', methods=['GET'])
@login_required
def export_examen_mental_pdf(exam_id):
    user_id = session.get('user_id')
    db = get_db()
    cursor = db.cursor()
    _ensure_examenes_mentales_table(cursor)
    
    cursor.execute("""
        SELECT e.*, p.nombres as pac_nombres, p.apellidos as pac_apellidos, p.cedula as pac_cedula,
               p.genero as pac_genero, p.fecha_nacimiento as pac_fecha_nacimiento, p.ocupacion as pac_ocupacion,
               u.nombres as psic_nombres, u.apellidos as psic_apellidos, u.estudios as psic_estudios, u.federacion as psic_federacion
        FROM examenes_mentales e
        JOIN pacientes p ON e.paciente_id = p.id
        JOIN usuarios u ON e.psicologo_id = u.id
        WHERE e.id = ? AND (e.psicologo_id = ? OR ? = 1)
    """, (exam_id, user_id, 1 if session.get('role') in ['admin', 'superadmin'] else 0))
    
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Examen mental no encontrado o sin permisos.'}), 404
        
    exam = dict(row)
    datos_eval = {}
    try:
        datos_eval = json.loads(exam['datos_evaluacion_json']) if exam.get('datos_evaluacion_json') else {}
    except: pass
    
    pac_nombre = f"{exam['pac_nombres']} {exam['pac_apellidos']}".strip()
    pac_cedula = exam['pac_cedula'] or 'Sin CI'
    psic_nombre = f"Psic. {exam['psic_nombres']} {exam['psic_apellidos']}".strip()
    psic_titulo = exam['psic_estudios'] or 'Psicólogo Clínico'
    psic_fed = exam['psic_federacion'] or 'S/N'
    
    # Renderizar vista HTML estandarizada lista para imprimir / PDF
    html = f"""
    <!DOCTYPE html>
    <html lang="es">
    <head>
        <meta charset="UTF-8">
        <title>Examen Mental - {pac_nombre}</title>
        <style>
            body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #f8fafc; color: #0f172a; padding: 20px; }}
            .card {{ max-width: 850px; margin: 0 auto; background: white; border-radius: 16px; padding: 2rem; box-shadow: 0 10px 30px rgba(0,0,0,0.08); }}
            .header {{ text-align: center; border-bottom: 2px solid #702e5e; padding-bottom: 1rem; margin-bottom: 1.5rem; }}
            .header h1 {{ margin: 0; color: #702e5e; font-size: 1.6rem; }}
            .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 1rem; background: #fdf4ff; border: 1px solid #f0abfc; padding: 1rem; border-radius: 10px; margin-bottom: 1.5rem; }}
            .section {{ margin-bottom: 1.2rem; padding: 1rem; background: #fafafa; border-left: 4px solid #702e5e; border-radius: 4px; }}
            .section h3 {{ margin: 0 0 0.5rem 0; color: #702e5e; font-size: 1.1rem; }}
            .no-print {{ text-align: center; margin-bottom: 15px; }}
            .btn-print {{ background: #702e5e; color: white; border: none; padding: 10px 20px; border-radius: 8px; font-weight: bold; cursor: pointer; }}
            @media print {{ .no-print {{ display: none; }} body {{ background: white; padding: 0; }} .card {{ box-shadow: none; padding: 0; }} }}
        </style>
    </head>
    <body>
        <div class="no-print">
            <button onclick="window.print()" class="btn-print">🖨️ Imprimir / Guardar en PDF</button>
        </div>
        <div class="card">
            <div class="header">
                <h1>INFORME DE EXAMEN DEL ESTADO MENTAL (MSE)</h1>
                <p>{psic_nombre} — {psic_titulo} (FPV: {psic_fed})</p>
            </div>
            <div class="grid">
                <div><strong>Consultante:</strong> {pac_nombre}</div>
                <div><strong>Cédula:</strong> {pac_cedula}</div>
                <div><strong>Fecha de Evaluación:</strong> {exam['fecha_evaluacion']}</div>
                <div><strong>Medio:</strong> {exam['medio_evaluacion']}</div>
            </div>
            
            {"".join([f"<div class='section'><h3>{cat.upper()}</h3><p>{json.dumps(val, ensure_ascii=False) if isinstance(val, (dict, list)) else val}</p></div>" for cat, val in datos_eval.items()])}
            
            <div class="section">
                <h3>OBSERVACIONES GENERALES Y CONCLUSIÓN</h3>
                <p>{exam['observaciones_generales'] or 'Sin observaciones adicionales.'}</p>
            </div>
            
            <div style="margin-top: 3rem; text-align: right;">
                <p>_________________________________________<br><strong>{psic_nombre}</strong><br>{psic_titulo} — FPV: {psic_fed}</p>
            </div>
        </div>
    </body>
    </html>
    """
    return render_template_string(html)

@examen_mental_bp.route('/api/examen-mental/<int:exam_id>/export/word', methods=['GET'])
@login_required
def export_examen_mental_word(exam_id):
    user_id = session.get('user_id')
    db = get_db()
    cursor = db.cursor()
    _ensure_examenes_mentales_table(cursor)
    
    cursor.execute("""
        SELECT e.*, p.nombres as pac_nombres, p.apellidos as pac_apellidos, p.cedula as pac_cedula,
               p.genero as pac_genero, p.fecha_nacimiento as pac_fecha_nacimiento, p.ocupacion as pac_ocupacion,
               u.nombres as psic_nombres, u.apellidos as psic_apellidos, u.estudios as psic_estudios, u.federacion as psic_federacion
        FROM examenes_mentales e
        JOIN pacientes p ON e.paciente_id = p.id
        JOIN usuarios u ON e.psicologo_id = u.id
        WHERE e.id = ? AND (e.psicologo_id = ? OR ? = 1)
    """, (exam_id, user_id, 1 if session.get('role') in ['admin', 'superadmin'] else 0))
    
    row = cursor.fetchone()
    if not row:
        return jsonify({'error': 'Examen mental no encontrado o sin permisos.'}), 404
        
    exam = dict(row)
    datos_eval = {}
    try:
        datos_eval = json.loads(exam['datos_evaluacion_json']) if exam.get('datos_evaluacion_json') else {}
    except: pass
    
    pac_nombre = f"{exam['pac_nombres']} {exam['pac_apellidos']}".strip()
    pac_cedula = exam['pac_cedula'] or 'Sin CI'
    psic_nombre = f"Psic. {exam['psic_nombres']} {exam['psic_apellidos']}".strip()
    psic_titulo = exam['psic_estudios'] or 'Psicólogo Clínico'
    psic_fed = exam['psic_federacion'] or 'S/N'
    
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("INFORME DE EXAMEN DEL ESTADO MENTAL (MSE)\n")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0x70, 0x2E, 0x5E)
        
        p_info = doc.add_paragraph()
        p_info.add_run(f"Consultante: {pac_nombre}\n").bold = True
        p_info.add_run(f"Cédula: {pac_cedula}\n")
        p_info.add_run(f"Fecha de Evaluación: {exam['fecha_evaluacion']}\n")
        p_info.add_run(f"Medio: {exam['medio_evaluacion']}\n")
        
        for cat, val in datos_eval.items():
            h = doc.add_heading(cat.upper(), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x70, 0x2E, 0x5E)
            doc.add_paragraph(json.dumps(val, ensure_ascii=False) if isinstance(val, (dict, list)) else str(val))
            
        doc.add_heading("OBSERVACIONES GENERALES", level=2)
        doc.add_paragraph(exam['observaciones_generales'] or 'Sin observaciones adicionales.')
        
        doc.add_paragraph("\n___________________________________________\n").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph(f"{psic_nombre}\n{psic_titulo} — FPV: {psic_fed}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        mem_file = io.BytesIO()
        doc.save(mem_file)
        mem_file.seek(0)
        
        clean_cedula = pac_cedula.replace(" ", "_")
        filename = f"Examen_Mental_{clean_cedula}_{exam['fecha_evaluacion']}.docx"
        
        return Response(
            mem_file.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({'error': f'Error al generar documento Word: {str(e)}'}), 500
