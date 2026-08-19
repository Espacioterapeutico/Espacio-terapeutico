import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Márgenes de página
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título Principal
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('Estructura y Diccionario de Datos\nMi Consultorio')
title_run.font.name = 'Arial'
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(93, 58, 111)

# Subtítulo
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Documentación Técnica de Tablas, Variables y Atributos de Base de Datos (SQLite)')
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(11)
sub_run.font.italic = True
sub_run.font.color.rgb = RGBColor(107, 114, 128)

doc.add_paragraph()

tables_data = [
    {
        'title': '1. usuarios (Psicólogos / Administradores)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID único del psicólogo (Llave primaria).'],
            ['username', 'TEXT', 'Nombre de usuario para iniciar sesión.'],
            ['password_hash', 'TEXT', 'Contraseña encriptada de forma segura.'],
            ['nombres', 'TEXT', 'Nombres del psicólogo.'],
            ['apellidos', 'TEXT', 'Apellidos del psicólogo.'],
            ['role', 'TEXT', 'Rol en el sistema (psicologo, admin).'],
            ['estudios', 'TEXT', 'Títulos o especialidades clínicas.'],
            ['federacion', 'TEXT', 'Número de colegiatura o federación.'],
            ['slug', 'TEXT', 'Identificador amigable para su enlace (ej: psic.paulomora).'],
            ['metodos_pago', 'TEXT', 'Cuentas bancarias / Pago móvil / Zelle / PayPal.'],
            ['disponibilidad_horarios', 'TEXT', 'Días y horas de consulta disponibles.'],
            ['terminos_condiciones', 'TEXT', 'Encuadre terapéutico personalizado.'],
            ['suscripcion_paga', 'INTEGER', '1 = Acceso Activo / 0 = Prueba.'],
            ['activo', 'INTEGER', '1 = Cuenta habilitada / 0 = Inactiva.']
        ]
    },
    {
        'title': '2. pacientes (Consultantes)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID único del consultante.'],
            ['psicologo_id', 'INTEGER', 'ID del psicólogo asignado.'],
            ['nombres', 'TEXT', 'Nombres del paciente.'],
            ['apellidos', 'TEXT', 'Apellidos del paciente.'],
            ['cedula', 'TEXT', 'Cédula / Documento de identidad.'],
            ['pronombre', 'TEXT', 'Pronombre de preferencia.'],
            ['genero', 'TEXT', 'Género del consultante.'],
            ['edad', 'INTEGER', 'Edad en años.'],
            ['fecha_nacimiento', 'TEXT', 'Fecha de nacimiento (YYYY-MM-DD).'],
            ['pais', 'TEXT', 'País de residencia.'],
            ['ciudad', 'TEXT', 'Ciudad de residencia.'],
            ['telefono', 'TEXT', 'Número de WhatsApp / Teléfono.'],
            ['email', 'TEXT', 'Correo electrónico.'],
            ['costo_personalizado', 'REAL', 'Tarifa especial por sesión (si aplica).'],
            ['moneda_personalizada', 'TEXT', 'Moneda acordada con el paciente (USD, EUR, etc.).'],
            ['costo_paquete_personalizado', 'REAL', 'Tarifa del paquete fraccionado (si aplica).'],
            ['sesiones_paquete_personalizado', 'INTEGER', 'Número de sesiones incluidas en paquete.'],
            ['terminos_aceptados', 'INTEGER', '1 = Aceptó encuadre terapéutico / 0 = Pendiente.'],
            ['username', 'TEXT', 'Usuario para acceder al portal PWA.'],
            ['password_hash', 'TEXT', 'Contraseña del paciente (Portal PWA).']
        ]
    },
    {
        'title': '3. agenda_finanzas (Consultas y Control Financiero)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID único del agendamiento.'],
            ['paciente_id', 'INTEGER', 'ID del consultante (Llave foránea).'],
            ['fecha', 'TEXT', 'Fecha de la consulta (YYYY-MM-DD).'],
            ['hora', 'TEXT', 'Hora de la consulta (HH:MM).'],
            ['tipo_consulta', 'TEXT', 'Modalidad (Presencial, Online, Uptaeb).'],
            ['monto', 'REAL', 'Precio acordado por la sesión.'],
            ['moneda', 'TEXT', 'Moneda (USD, EUR, BSD, COP, ARS).'],
            ['estado_pago', 'TEXT', 'Estado (Paga, Pendiente, Prepagada, Exonerado).'],
            ['control_uso', 'TEXT', 'Consumida o No consumida.'],
            ['confirmada', 'INTEGER', '1 = Confirmada por WhatsApp / 0 = Por confirmar.'],
            ['recordatorio_enviado_wa', 'INTEGER', '1 = Recordatorio enviado / 0 = Pendiente.'],
            ['confirmacion_enviada_wa', 'INTEGER', '1 = Solicitud enviada / 0 = Pendiente.'],
            ['metodo_pago', 'TEXT', 'Forma de pago utilizada (ej: Pago Móvil, Zelle).'],
            ['referencia', 'TEXT', 'Número de referencia del comprobante.']
        ]
    },
    {
        'title': '4. sesiones (Evoluciones y Notas Clínicas)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID de la nota evolutiva.'],
            ['paciente_id', 'INTEGER', 'ID del consultante.'],
            ['agenda_id', 'INTEGER', 'ID de la cita en agenda_finanzas.'],
            ['fecha', 'TEXT', 'Fecha de la sesión.'],
            ['modalidad', 'TEXT', 'Modalidad empleada.'],
            ['estado', 'TEXT', 'Realizada, Cancelada, Reprogramada.'],
            ['resumen', 'TEXT', 'Evolución clínica privada del psicólogo.'],
            ['resumen_paciente', 'TEXT', 'Resumen visible para el paciente en su portal.'],
            ['tareas_asignadas', 'TEXT', 'Tareas / compromisos asignados al paciente.'],
            ['recursos_entregados', 'TEXT', 'Herramientas o guías facilitadas.'],
            ['anotaciones_proxima', 'TEXT', 'Puntos pendientes para repasar próxima cita.'],
            ['diagnostico', 'TEXT', 'Diagnóstico o hipótesis clínica.'],
            ['archivo_adjunto', 'TEXT', 'Ruta de imágenes/PDFs subidos en sesión.']
        ]
    },
    {
        'title': '5. tarifas_pais (Arancel Multimoneda por País)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID único del arancel.'],
            ['psicologo_id', 'INTEGER', 'ID del psicólogo.'],
            ['pais', 'TEXT', 'Nombre del país (ej: Venezuela, Colombia).'],
            ['modalidad', 'TEXT', 'Online o Presencial.'],
            ['costo_individual', 'REAL', 'Precio de la sesión suelta.'],
            ['costo_paquete', 'REAL', 'Precio del paquete completo.'],
            ['sesiones_paquete', 'INTEGER', 'Cantidad de sesiones del paquete.'],
            ['moneda', 'TEXT', 'Moneda aplicable (USD, COP, EUR, etc.).']
        ]
    },
    {
        'title': '6. pagos_notificados (Comprobantes Reportados por Pacientes)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID del comprobante reportado.'],
            ['paciente_id', 'INTEGER', 'ID del paciente.'],
            ['monto', 'REAL', 'Monto transferido.'],
            ['moneda', 'TEXT', 'Moneda del pago.'],
            ['metodo', 'TEXT', 'Método (Pago Móvil, Zelle, etc.).'],
            ['referencia', 'TEXT', 'Número de referencia bancaria.'],
            ['estado', 'TEXT', 'Pendiente de verificación, Aprobado, Rechazado.']
        ]
    },
    {
        'title': '7. pizarra_terapeutica (Diario y Estado de Ánimo del Paciente)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['id', 'INTEGER', 'ID del registro.'],
            ['paciente_id', 'INTEGER', 'ID del paciente.'],
            ['fecha', 'TEXT', 'Fecha del registro.'],
            ['estado_animo', 'TEXT', 'Emoción (ej: Feliz, Ansioso).'],
            ['emoji_animo', 'TEXT', 'Emoji representativo (ej: 😊, 😰).'],
            ['contenido', 'TEXT', 'Escrito reflexivo del paciente.']
        ]
    },
    {
        'title': '8. configuracion (Parámetros Globales y Plantillas)',
        'cols': ['Campo (Variable)', 'Tipo', 'Descripción'],
        'rows': [
            ['clave', 'TEXT', 'Identificador de configuración (msg_confirmacion, msg_recordatorio, etc.).'],
            ['valor', 'TEXT', 'Contenido o texto de la plantilla/parámetro.']
        ]
    }
]

for tspec in tables_data:
    h = doc.add_heading(tspec['title'], level=2)
    h.runs[0].font.name = 'Arial'
    h.runs[0].font.size = Pt(14)
    h.runs[0].font.color.rgb = RGBColor(93, 58, 111)
    
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    
    hdr_cells = t.rows[0].cells
    for i, col_name in enumerate(tspec['cols']):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="5D3A6F"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        
    for rdata in tspec['rows']:
        row_cells = t.add_row().cells
        for i, val in enumerate(rdata):
            row_cells[i].text = val
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(9.5)
            if i == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(169, 89, 147)
                
    doc.add_paragraph()

output_filename = 'Estructura_Base_de_Datos_Mi_Consultorio.docx'
doc.save(output_filename)
print(f'Documento Word generado exitosamente: {output_filename}')
